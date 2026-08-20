#!/usr/bin/env python3
"""
批量生成33个采砂场监测报告并合并为一个文件。

每份报告依次调用 caisha_report_tool，中间失败跳过继续。
合并时按《2025采砂报告》章节结构输出编号层级：
  1.批复采砂区年度实施情况监测与评估
  1.X县区采砂区监测与评估
  1.X.Y站点标题（采砂场=许可证号 / 工程=规范工程名）
  1.X.Y.1 基本情况  1.X.Y.2 监测实施情况
  （每县末尾追加"1.X.(Y+1)县区采砂场监测评估意见"占位标题，意见文字由人工填写）

merge_reports() 可独立导入使用（单元测试见 tests/test_batch_merge.py）。
"""
import copy
import json
import logging
import os
import sys
import time

BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BACKEND_DIR)
os.chdir(BACKEND_DIR)

# 日志配置（同时输出到控制台和文件）
log_path = os.path.join(BACKEND_DIR, 'logs', 'batch_caisha_reports.log')
os.makedirs(os.path.dirname(log_path), exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(log_path, encoding='utf-8'),
    ])
logger = logging.getLogger('batch_caisha')

# 加载 .env（DASHSCOPE_API_KEY / PROJ_LIB 等）
env_path = os.path.join(BACKEND_DIR, '.env')
if os.path.exists(env_path):
    with open(env_path, encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith('#') or '=' not in line:
                continue
            k, v = line.split('=', 1)
            os.environ.setdefault(k.strip(), v.strip())
    logger.info('已加载 .env (%d 行)', sum(1 for _ in open(env_path)))

# 修复 PROJ_LIB 版本不兼容（必须用 rasterio 自带的 proj.db）
_rasterio_proj = '/home/server/miniconda3/envs/mapagent6/lib/python3.11/site-packages/rasterio/proj_data'
if os.path.isdir(_rasterio_proj):
    os.environ['PROJ_LIB'] = _rasterio_proj
    os.environ['PROJ_DATA'] = _rasterio_proj  # pyproj ≥3.1 用这个变量
    logger.info('PROJ_LIB/PROJ_DATA=%s', _rasterio_proj)
else:
    logger.warning('rasterio proj_data 目录不存在: %s，可能影响 TIF 读取', _rasterio_proj)

from tools.caisha_report_tool import CaishaReportTool  # noqa: E402

# 每个县区末尾追加"采砂场监测评估意见"占位标题（意见文字人工填写）
ADD_OPINION_PLACEHOLDER = True

# 33个采砂场（序号≤34，按文件夹编号排序；第三列为县区，用于合并章节编号）
SITES = [
    ("淮河郑湾村可采区",               "1豫信平砂许[2025]第1号", "平桥区"),
    ("高店镇王湾村部分河段",             "2豫罗砂许[2025]第01号", "罗山县"),
    ("浉淮村采区",                     "3豫罗砂许[2025]第02号", "罗山县"),
    ("山店乡青莲村部分河段",             "4豫罗砂许[2025]第03号", "罗山县"),
    ("青龙闵塆采区",                    "5豫罗砂许[2025]第04号", "罗山县"),
    ("中山村采区",                     "6豫罗砂许[2025]第05号", "罗山县"),
    ("淮河村采区",                     "7豫罗砂许[2025]第06号", "罗山县"),
    ("天湖郑洼采区",                   "8豫罗砂许[2025]第07号", "罗山县"),
    ("吴乡村采区",                     "9豫罗砂许[2025]第08号", "罗山县"),
    ("楠杆镇李寨村、邵湾村部分河段",       "11豫罗砂许[2025]第10号", "罗山县"),
    ("竹竿河息县八里岔乡李庄可采区",       "12豫息水砂许[2025]第01号", "息县"),
    ("郝楼砂场",                       "13豫潢砂许[2025]第1号", "潢川县"),
    ("黄寨砂场",                       "14豫潢砂许[2025]第2号", "潢川县"),
    ("石牛村采区",                     "15豫商城砂许[2025]第01号", "商城县"),
    ("金寨村采区",                     "16豫商城砂许[2025]第02号", "商城县"),
    ("王楼村采区",                     "17豫商城砂许[2025]第03号", "商城县"),
    ("李湖村采区",                     "18豫商城砂许[2025]第04号", "商城县"),
    ("史灌河童庙可采区",                 "19豫信固砂许[2025]第01号", "固始县"),
    ("史灌河大营-范营可采区",            "20豫信固砂许[2025]第02号", "固始县"),
    ("史河汪营可采区",                  "21豫信固砂许[2025]第03号", "固始县"),
    ("史河红石可采区",                  "22豫信固砂许[2025]第04号", "固始县"),
    ("陈营-柳沟采区",                  "23豫信固砂许[2025]第05号", "固始县"),
    ("郑堂-大庄采区",                  "24豫信固砂许[2025]第06号", "固始县"),
    ("史河余庆-人主可采区",             "25豫信固砂许[2025]第07号", "固始县"),
    ("南元、祝家楼采区",                "26豫信固砂许[2025]第08号", "固始县"),
    ("牛老家、龙潭采区",                "27豫信固砂许[2025]第09号", "固始县"),
    ("史河毕店可采区",                  "28豫信固砂许[2025]第10号", "固始县"),
    ("史河泗湖-长兴可采区",             "29豫信固砂许[2025]第11号", "固始县"),
    ("灌河淮堰-周集可采区",             "30豫信固砂许[2025]第12号", "固始县"),
    ("灌河道超可采区",                  "31豫信固砂许[2025]第13号", "固始县"),
    ("灌河叶岗可采区",                  "33豫信固砂许[2025]第15号", "固始县"),
    ("沈营可采区",                      "34豫信固砂许[2025]第16号", "固始县"),
]

# 县区章节顺序（与《2025采砂报告》一致）
COUNTY_ORDER = ['平桥区', '罗山县', '息县', '潢川县', '商城县', '固始县']


def merge_reports(success_files, out_dir, add_opinion=True):
    """合并已生成报告：success_files = [(filepath, site, county, variables), ...]。
    保留图片，插入县区章节与三级编号（1.X.Y 式），返回合并文件路径。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.parts.image import ImagePart
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Pt

    R_EMBED = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    R_ID = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'

    def _set_cn_font(run, size=12, bold=False):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

    def add_para(doc, text, size=12, bold=False, align=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        _set_cn_font(r, size=size, bold=bold)
        if align is not None:
            p.alignment = align
        p.paragraph_format.line_spacing = 1.5
        return p

    def rewrite_para(p, text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        run = p.add_run(text)
        _set_cn_font(run, size=size, bold=bold)
        if align is not None:
            p.alignment = align

    # 以第一份成功报告为母版（保留其样式资源），清空正文后按编号结构重排
    merged = Document(success_files[0][0])
    for el in list(merged.element.body):
        merged.element.body.remove(el)

    add_para(merged, '1.批复采砂区年度实施情况监测与评估', size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    county_no = {c: i + 1 for i, c in enumerate(COUNTY_ORDER)}
    county_site_count = {}   # county -> 已合并站点数
    pending_opinions = []    # (county, 章节号, 站点数)

    for filepath, site, county, variables in success_files:
        src = Document(filepath)
        c_no = county_no.get(county, 0)
        if county not in county_site_count:
            # 上一县区末尾补"采砂场监测评估意见"占位
            if pending_opinions and add_opinion:
                for pc, pc_no, pn in pending_opinions:
                    add_para(merged, f'1.{pc_no}.{pn + 1}{pc}采砂场监测评估意见',
                             size=14, bold=True)
                    add_para(merged, '', size=12)
            pending_opinions = []
            county_site_count[county] = 0
            add_para(merged, '')  # 空行
            add_para(merged, f'1.{c_no}{county}采砂区监测与评估', size=15, bold=True)
        county_site_count[county] += 1
        s_no = county_site_count[county]

        site_title = variables.get('site_title') or site
        basic_sec = variables.get('basic_section_name') or '采砂场基本情况'
        monitor_sec = variables.get('monitor_section_name') or '采砂场监测实施情况'

        # 复制图片关系到合并文档
        rid_map = {}
        for rId, rel in src.part.rels.items():
            reltype_str = str(rel.reltype)
            if 'image' in reltype_str:
                try:
                    blob = rel.target_part.blob
                    content_type = rel.target_part.content_type
                    partname = merged.part.package.next_partname('/word/media/image%d.png')
                    img_part = ImagePart(partname, content_type, blob)
                    new_rId = merged.part.relate_to(img_part, RT.IMAGE)
                    rid_map[rId] = new_rId
                except Exception as e:
                    logger.warning('图片复制失败 %s: %s', filepath, e)

        # 逐元素克隆并修正图片引用（跳过 sectPr：多个节属性会导致后续
        # add_paragraph 全部插入到第一个 sectPr 之前，章节/意见占位错位）
        start_para = len(merged.paragraphs)
        for elem in src.element.body:
            if elem.tag.endswith('}sectPr'):
                continue
            cloned = copy.deepcopy(elem)
            for el in cloned.iter():
                for attr_name in (R_EMBED, R_ID):
                    old = el.attrib.get(attr_name)
                    if old and old in rid_map:
                        el.attrib[attr_name] = rid_map[old]
            merged.element.body.append(cloned)

        # 改写本站点标题与两级小节编号（只处理刚追加的段落）
        for p in merged.paragraphs[start_para:]:
            t = p.text.strip()
            if t == site_title:
                rewrite_para(p, f'1.{c_no}.{s_no}{site_title}', size=14, bold=True)
            elif t == f'一、{basic_sec}':
                rewrite_para(p, f'1.{c_no}.{s_no}.1 {basic_sec}', size=14, bold=True)
            elif t == f'二、{monitor_sec}':
                rewrite_para(p, f'1.{c_no}.{s_no}.2{monitor_sec}', size=14, bold=True)

        pending_opinions.append((county, c_no, s_no))

    # 最后一个县区的评估意见占位
    if pending_opinions and add_opinion:
        pc, pc_no, pn = pending_opinions[-1]
        add_para(merged, f'1.{pc_no}.{pn + 1}{pc}采砂场监测评估意见', size=14, bold=True)
        add_para(merged, '', size=12)

    os.makedirs(out_dir, exist_ok=True)
    merged_path = os.path.join(out_dir, f'caisha_merged_{int(time.time())}.docx')
    merged.save(merged_path)
    return merged_path


def main():
    tool = CaishaReportTool()
    out_dir = os.path.join(BACKEND_DIR, 'static', 'reports')
    os.makedirs(out_dir, exist_ok=True)

    success_files = []  # (filepath, site, county, variables)
    fail_sites = []

    t_total = time.time()
    for idx, (site, folder_hint, county) in enumerate(SITES, 1):
        print(f'\n{"="*60}')
        print(f'[{idx}/{len(SITES)}] {site}')
        print(f'{"="*60}')
        t0 = time.time()
        try:
            result = tool.call({
                'site_name': site,
                # permit_no 从数据库自动读取，不传即可
            })
            if result.get('success'):
                filepath = result.get('file_path', '')
                if filepath and os.path.exists(filepath):
                    success_files.append((filepath, site, county, result.get('variables') or {}))
                    elapsed = time.time() - t0
                    fsize = os.path.getsize(filepath)
                    notes_str = str(result.get('notes', ''))
                    has_ortho = '影像图缺失' not in notes_str and '影像裁剪失败' not in notes_str
                    status_parts = [f'{elapsed:.0f}s', f'{fsize/1024:.0f}KB']
                    if has_ortho:
                        status_parts.append('正射✓')
                    else:
                        status_parts.append('正射✗')
                    if '失败' not in notes_str:
                        status_parts.append('高程✓')
                    title = result.get('variables', {}).get('site_title', site)
                    print(f'  OK: {os.path.basename(filepath)}  ({" | ".join(status_parts)}) 标题={title}')
                else:
                    fail_sites.append((site, '文件路径无效'))
                    print(f'  FAIL: 文件路径无效')
            else:
                err = result.get('error', 'unknown')
                fail_sites.append((site, err))
                print(f'  FAIL: {err}')
        except Exception as e:
            fail_sites.append((site, str(e)))
            print(f'  EXCEPTION: {e}')

    elapsed_total = time.time() - t_total
    print(f'\n{"="*60}')
    print(f'完成：成功 {len(success_files)} / 失败 {len(fail_sites)} / 总计 {len(SITES)}')
    print(f'总耗时 {elapsed_total:.0f}s ({elapsed_total/60:.1f}min)')
    if fail_sites:
        print('\n失败列表:')
        for s, e in fail_sites:
            print(f'  - {s}: {e}')

    # ── 合并 docx ──
    if len(success_files) >= 2:
        print(f'\n正在合并 {len(success_files)} 份报告（保留图片 + 编号章节）...')
        merged_path = merge_reports(success_files, out_dir, ADD_OPINION_PLACEHOLDER)
        print(f'合并完成: {merged_path}  ({os.path.getsize(merged_path)/1024:.0f}KB)')
    elif len(success_files) == 1:
        print(f'只有1份成功，无需合并: {success_files[0][0]}')
    else:
        print('所有报告均失败，无法合并')


if __name__ == '__main__':
    main()
