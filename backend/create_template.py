from docx import Document
import os

def create_template():
    doc = Document()
    doc.add_heading('{{report_title}}', 0)

    doc.add_heading('1. 概述', level=1)
    doc.add_paragraph('{{summary}}')

    doc.add_heading('2. 详细数据', level=1)
    doc.add_paragraph('根据您的查询，以下是相关数据分析：')
    doc.add_paragraph('{{details}}')

    doc.add_heading('3. 结论', level=1)
    doc.add_paragraph('{{conclusion}}')

    doc.add_paragraph('生成时间: {{generated_date}}')

    output_path = '/home/server/python/map_assistant_v1/backend/templates/report_template.docx'
    doc.save(output_path)
    print(f"Template created at {output_path}")

if __name__ == "__main__":
    create_template()
