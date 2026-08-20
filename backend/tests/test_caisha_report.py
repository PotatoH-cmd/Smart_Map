# -*- coding: utf-8 -*-
"""
caisha_report_tool 纯逻辑单元测试（不依赖 DB / 网络 / LLM API）。

覆盖 2026-08 批次改造：
- 站点类型判定（采砂场 vs 疏浚工程）
- 规范站点名（工程类去"南-南1/-01"分段后缀）
- 许可证号规范化（括号统一〔〕、去空格）
- 基本情况散文生成的兜底链路（无 KB / LLM 失败）
- 模板渲染：无残留占位符、标题/图注使用规范化名称、图号顺延
"""
import os
import sys

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from tools.caisha_report_tool import BASIC_FIELDS, MISSING, CaishaReportTool  # noqa: E402


@pytest.fixture(scope='module')
def tool():
    return CaishaReportTool()


class TestSiteClassify:
    def test_project_detection(self, tool):
        assert tool._is_project('潢川县白露河上游防洪能力提升工程南-南1') is True
        assert tool._is_project('淮河(竹竿河入淮口至桃花岛段)航道疏浚及岸线整治工程') is True
        assert tool._is_project('固始县马家河生态修复项目') is True
        assert tool._is_project('淮河闾河口港区港池疏浚工程') is True
        assert tool._is_project('固始县2021年重点水毁修复黎集滚水坝上游清淤扩容工程') is True

    def test_sand_detection(self, tool):
        assert tool._is_project('郝楼砂场') is False
        assert tool._is_project('青龙闵塆采区') is False
        assert tool._is_project('高店镇王湾村部分河段') is False
        assert tool._is_project('史河毕店可采区') is False


class TestDisplayName:
    def test_project_segment_suffix_stripped(self, tool):
        assert tool._display_name('潢川县白露河上游防洪能力提升工程南-南1', True) == \
            '潢川县白露河上游防洪能力提升工程'
        assert tool._display_name('潢川县白露河上游防洪能力提升工程-01', True) == \
            '潢川县白露河上游防洪能力提升工程'

    def test_project_name_kept(self, tool):
        assert tool._display_name('淮河(竹竿河入淮口至桃花岛段)航道疏浚及岸线整治工程', True) == \
            '淮河(竹竿河入淮口至桃花岛段)航道疏浚及岸线整治工程'

    def test_sand_name_kept(self, tool):
        assert tool._display_name('郝楼砂场', False) == '郝楼砂场'
        assert tool._display_name('史河毕店可采区', False) == '史河毕店可采区'

    def test_rtk_stripped(self, tool):
        assert tool._display_name('郝楼砂场RTK', False) == '郝楼砂场'


class TestPermitNormalize:
    def test_brackets_unified(self, tool):
        assert tool._normalize_permit('信水河[2024]15号') == '信水河〔2024〕15号'
        assert tool._normalize_permit('豫潢砂许（2025）第1号') == '豫潢砂许〔2025〕第1号'
        assert tool._normalize_permit('豫潢砂许〔2025〕第1号') == '豫潢砂许〔2025〕第1号'

    def test_spaces_removed(self, tool):
        assert tool._normalize_permit(' 豫潢砂许 〔2025〕 第 1 号 ') == '豫潢砂许〔2025〕第1号'

    def test_missing(self, tool):
        assert tool._normalize_permit('') == MISSING
        assert tool._normalize_permit(MISSING) == MISSING


class TestBasicProse:
    @pytest.fixture
    def sand_fields(self):
        return {k: MISSING for k, _ in BASIC_FIELDS}

    def test_project_no_kb_fallback(self, tool, monkeypatch):
        monkeypatch.setattr(tool, '_retrieve_basic_chunks', lambda area: ([], []))
        prose, src = tool._extract_basic_prose(
            '潢川县白露河上游防洪能力提升工程', '潢川县白露河上游防洪能力提升工程南-南1',
            True, {'location': MISSING})
        assert '潢川县白露河上游防洪能力提升工程' in prose
        assert MISSING in prose
        assert src == ''

    def test_sand_no_kb_fallback(self, tool, monkeypatch, sand_fields):
        monkeypatch.setattr(tool, '_retrieve_basic_chunks', lambda area: ([], []))
        prose, src = tool._extract_basic_prose('郝楼村采区', '郝楼砂场', False, sand_fields)
        assert '郝楼砂场采砂区域位于' in prose
        assert prose.count(MISSING) >= 5
        assert src == ''

    def test_llm_failure_falls_back(self, tool, monkeypatch, sand_fields):
        monkeypatch.setattr(tool, '_retrieve_basic_chunks', lambda area: (['【x】原文片段'], ['x']))
        def _boom(_p):
            raise RuntimeError('no api')
        monkeypatch.setattr(tool, '_llm', _boom)
        prose, src = tool._extract_basic_prose('郝楼村采区', '郝楼砂场', False, sand_fields)
        assert '郝楼砂场采砂区域位于' in prose
        assert src == ''

    def test_bad_llm_output_falls_back(self, tool, monkeypatch, sand_fields):
        monkeypatch.setattr(tool, '_retrieve_basic_chunks', lambda area: (['【x】原文片段'], ['x']))
        monkeypatch.setattr(tool, '_llm', lambda _p: chr(96) * 3 + 'json\n{}')
        prose, _ = tool._extract_basic_prose('郝楼村采区', '郝楼砂场', False, sand_fields)
        assert '郝楼砂场采砂区域位于' in prose

    def test_valid_llm_prose_returned(self, tool, monkeypatch, sand_fields):
        monkeypatch.setattr(tool, '_retrieve_basic_chunks', lambda area: (['【x】原文片段'], ['x']))
        monkeypatch.setattr(tool, '_llm',
                            lambda _p: '郝楼砂场采砂区域位于潢河魏岗镇郝楼村河段，开采企业为河南水投光州现代水网有限公司。')
        prose, src = tool._extract_basic_prose('郝楼村采区', '郝楼砂场', False, sand_fields)
        assert prose.startswith('郝楼砂场采砂区域位于')
        assert src == 'x'


class TestRenderDocx:
    def _variables(self):
        v = {
            'site_name': '郝楼砂场',
            'site_title': '豫潢砂许〔2025〕第1号',
            'display_name': '郝楼砂场',
            'permit_no': '豫潢砂许〔2025〕第1号',
            'plan_year': '2025',
            'basic_info': '郝楼砂场采砂区域位于潢河魏岗镇郝楼村河段，开采企业为河南水投光州现代水网有限公司。',
            'basic_section_name': '采砂场基本情况',
            'monitor_section_name': '采砂场监测实施情况',
            'drone_subject': '豫潢砂许〔2025〕第1号采砂区域',
            'drone_scope': '2025年度规划批复范围（桩号K11+900～K12+721）',
            'field_monitoring': '2026年6月对郝楼砂场开展现场监测，发现该砂场处于停工状态。',
            'drone_interpretation': '经解译，采区内无明显作业迹象，处于停工状态。',
            'depth_evaluation': '实测平均高程25.019m，不低于控制开采高程下限22.19m，本年度开采深度基本符合年度实施方案。',
            'ctrl_bottom_elevation': '22.41～22.19m',
            'point_count': '7881',
            'elev_min': '19.150',
            'elev_max': '30.910',
            'elev_avg': '25.019',
            'stake_range': 'K11+900～K12+721',
            'fig_ortho_no': '3',
            'fig_elev_no': '4',
            'generated_date': '2026年08月19日',
        }
        for k, _ in BASIC_FIELDS:
            v.setdefault(k, MISSING)
        return v

    def test_render_no_placeholder_left(self, tool, tmp_path):
        from PIL import Image
        ortho = tmp_path / 'o.png'
        Image.new('RGB', (80, 80), 'blue').save(ortho)
        elev = tmp_path / 'e.png'
        Image.new('RGB', (80, 80), 'red').save(elev)

        result = tool._render_docx(self._variables(), str(ortho), str(elev), [str(ortho), str(elev)])
        assert result.get('success'), result
        assert os.path.exists(result['file_path'])

        from docx import Document
        doc = Document(result['file_path'])
        full = chr(10).join(p.text for p in doc.paragraphs)
        assert '{{' not in full, '模板占位符残留'
        assert '豫潢砂许〔2025〕第1号' in full            # 标题=许可证号
        assert '一、采砂场基本情况' in full
        assert '图1 郝楼砂场现场监测图片（一）' in full     # 图注用规范站点名 + 图号（一）
        assert '图2 郝楼砂场现场监测图片（二）' in full
        assert '图3 郝楼砂场无人机正射影像图' in full      # 图号顺延
        assert '图4 郝楼砂场高程点示意图' in full
        os.remove(result['file_path'])

    def test_render_project_title(self, tool, tmp_path):
        from PIL import Image
        png = tmp_path / 'o.png'
        Image.new('RGB', (80, 80), 'blue').save(png)
        v = self._variables()
        v.update({
            'site_name': '潢川县白露河上游防洪能力提升工程南-南1',
            'site_title': '潢川县白露河上游防洪能力提升工程',
            'display_name': '潢川县白露河上游防洪能力提升工程',
            'basic_section_name': '疏浚工程基本情况',
            'monitor_section_name': '疏浚工程监测实施情况',
            'drone_subject': '潢川县白露河上游防洪能力提升工程',
            'drone_scope': '批复工程范围',
        })
        result = tool._render_docx(v, str(png), str(png), [])
        assert result.get('success'), result
        from docx import Document
        doc = Document(result['file_path'])
        full = chr(10).join(p.text for p in doc.paragraphs)
        assert '潢川县白露河上游防洪能力提升工程' in full
        assert '一、疏浚工程基本情况' in full
        assert '叠加批复工程范围进行评估分析' in full
        assert '（许可证号' not in full, '工程类标题不再带许可证号行'
        os.remove(result['file_path'])
