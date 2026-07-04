#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修改报告模板，添加知识库内容占位符
"""

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

def add_knowledge_placeholder():
    doc = Document('/home/server/python/map_assistant_v1/backend/templates/report_template.docx')
    
    # 查找 conclusion 占位符的位置
    conclusion_idx = None
    for i, para in enumerate(doc.paragraphs):
        if '{{conclusion}}' in para.text:
            conclusion_idx = i
            break
    
    if conclusion_idx is None:
        print('未找到 conclusion 占位符，尝试查找 summary')
        for i, para in enumerate(doc.paragraphs):
            if '{{summary}}' in para.text:
                conclusion_idx = i
                break
    
    if conclusion_idx is not None:
        print(f'在段落 {conclusion_idx} 找到目标占位符')
        
        # 获取原始段落样式
        ref_para = doc.paragraphs[conclusion_idx]
        ref_style = ref_para.style
        
        # 创建新的段落并插入知识库占位符
        # 在 conclusion 前面插入
        new_para = doc.paragraphs[conclusion_idx].insert_paragraph_before()
        new_para.text = ''
        run = new_para.add_run()
        run.text = '三、政策依据与参考资料'
        
        # 添加知识库内容段落
        kb_para = doc.paragraphs[conclusion_idx].insert_paragraph_before()
        kb_para.text = ''
        run = kb_para.add_run()
        run.text = '{{knowledge_content}}'
        
        print('已添加知识库内容段落')
        
        # 保存
        doc.save('/home/server/python/map_assistant_v1/backend/templates/report_template.docx')
        print('模板已保存')
        
        return True
    else:
        print('未找到任何目标占位符')
        return False

if __name__ == '__main__':
    add_knowledge_placeholder()
