# -*- coding: utf-8 -*-
"""
生成采砂场监测报告 docx 模板：templates/caisha_monitor_template.docx

结构对照《2025采砂报告》采砂场监测章节（1.X.Y 式，单站报告不带编号）：
  {{site_title}}                       —— 标题：采砂场=许可证号（豫潢砂许〔2025〕第1号），工程=规范工程名
  一、{{basic_section_name}}           —— 采砂场基本情况 / 疏浚工程基本情况
  {{basic_info}}                       —— 知识库散文段（按参考文风，缺失写"资料未见"）
  二、{{monitor_section_name}}         —— 采砂场监测实施情况 / 疏浚工程监测实施情况
  （一）现场监测 {{field_monitoring}} + 图1..N 现场监测图片
  （二）无人机航测 {{drone_subject}} + {{drone_scope}} + 正射影像图
  （三）无人船水下测量 + 高程点示意图
图片占位段落 {{img_field_photos}} / {{img_drone_ortho}} / {{img_elev_points}}
由 caisha_report_tool 原位替换为插图；现场监测图片数量不定，图号顺延，
故正射/高程图的图号用 {{fig_ortho_no}} / {{fig_elev_no}} 动态填充。
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def _set_cn_font(run, size=12, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')


def _para(doc, text, size=12, indent=True, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn_font(run, size=size, bold=bold)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    return p


def create_template():
    doc = Document()

    # 标题（采砂场=许可证号；工程=规范工程名；由 caisha_report_tool 计算）
    title = doc.add_paragraph()
    run = title.add_run('{{site_title}}')
    _set_cn_font(run, size=18, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============ 一、基本情况（类型感知标题 + 知识库散文） ============
    _para(doc, '一、{{basic_section_name}}', size=15, indent=False, bold=True)
    _para(doc, '{{basic_info}}')

    # ============ 二、监测实施情况 ============
    _para(doc, '二、{{monitor_section_name}}', size=15, indent=False, bold=True)

    # （一）现场监测
    _para(doc, '（一）现场监测', size=14, indent=False, bold=True)
    _para(doc, '{{field_monitoring}}')
    _para(doc, '{{img_field_photos}}', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    # （二）无人机航测（采砂场：对许可证号采砂区域+桩号；工程：对工程名+批复工程范围）
    _para(doc, '（二）无人机航测', size=14, indent=False, bold=True)
    _para(doc,
          '采砂监管测量人员对{{drone_subject}}进行了无人机航空摄影测量，'
          '叠加{{drone_scope}}进行评估分析，{{drone_interpretation}}')
    _para(doc, '{{img_drone_ortho}}', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, '图{{fig_ortho_no}} {{display_name}}无人机正射影像图', size=10.5, indent=False,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    # （三）无人船水下测量
    _para(doc, '（三）无人船水下测量', size=14, indent=False, bold=True)
    _para(doc,
          '根据批复文件和{{plan_year}}年度实施方案，开采后控制高程为{{ctrl_bottom_elevation}}，'
          '通过无人船单波束声呐测量采区水下地形，对采区开采深度进行评估分析。'
          '实测高程点共{{point_count}}个，高程范围{{elev_min}}～{{elev_max}}m，'
          '平均高程{{elev_avg}}m。{{depth_evaluation}}')
    _para(doc, '{{img_elev_points}}', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, '图{{fig_elev_no}} {{display_name}}高程点示意图', size=10.5, indent=False,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    # 落款
    _para(doc, '报告生成时间：{{generated_date}}', size=10.5, indent=False,
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    out_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates')
    os.makedirs(out_dir, exist_ok=True)
    output_path = os.path.join(out_dir, 'caisha_monitor_template.docx')
    doc.save(output_path)
    print(f'模板已生成: {output_path}')
    return output_path


if __name__ == '__main__':
    create_template()
