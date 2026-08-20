# -*- coding: utf-8 -*-
"""
采砂场监测报告生成工具（caisha_report_tool）

输入砂场名/许可证号，五步生成监测分析报告：
  1. 基本情况：LlamaIndex 实施方案文件夹检索 + LLM 抽取 11 个变量（严格取原文，缺失填"资料未见"）
  2. 现场监测：现场监测报告文件夹检索整理
  3. 无人机解译：按采区范围裁剪正射影像 + 叠加批复边界，送 qwen-vl 判读（停工/作业器具/堆砂堆靠）
  4. 无人船评估：rtk_points 平均高程 vs 控制开采高程 → 三档结论
  5. 高程点示意图：影像底图 + 采区边线 + 红色测点高程标注（同图3.5-74）

最后填充 templates/caisha_monitor_template.docx 生成报告。
"""
import json
import logging
import os
import re
import time
from typing import Any, Dict, List, Optional, Tuple, Union

from qwen_agent.tools.base import BaseTool, register_tool

logger = logging.getLogger(__name__)

BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

DB_CFG = {
    "host": "172.136.16.52",
    "port": 5432,
    "dbname": "postgres",
    "user": "postgres",
    "password": "8720622",
}

# 批复采区边界（EPSG:4326，Name 字段为砂场名）
CAIQU_SHP = "/home/server/python/map_assistant_v1/data/2026年采区.shp"
# 河道管理范围线（线要素，左岸/右岸分开，与前端地图"红线底图"同源；KML 监管范围为面要素，
# 画出来两岸会连成闭合环，故不用 KML）
HEDAO_GEOJSON = "/home/server/python/map_assistant_v1/frontend/public/data/hx.geojson"
# 默认无人机正射影像
DEFAULT_TIF = "/home/server/python/采砂/潢川县.tif"
# 砂场专属彩色正射影像目录（文件名=砂场名.tif，RGB 彩色，优先于全县灰度大图）
SITE_TIF_DIR = "/home/server/python/采砂/砂场影像"

# CGCS2000 3度带显式 Proj4 参数（去带号假东 x_0=500000）：
# - 环境 PROJ 库版本混乱（share/proj 旧 layout 被 GDAL 拒绝，pyproj 自带库定义与官方
#   不一致），from_epsg 查库结果不确定，必须用显式参数；
# - 无人机影像与 RTK 数据坐标均为去带号米制（x≈550xxx），标准带号假东
#   （38 带=38500000 / 39 带=39500000）会错开一个带号，投影边界与影像永远无交集。
FALLBACK_PROJ4 = {
    4547: '+proj=tmerc +lat_0=0 +lon_0=114 +k=1 +x_0=500000 +y_0=0 +ellps=GRS80 +units=m +no_defs',
    4548: '+proj=tmerc +lat_0=0 +lon_0=117 +k=1 +x_0=500000 +y_0=0 +ellps=GRS80 +units=m +no_defs',
    4549: '+proj=tmerc +lat_0=0 +lon_0=120 +k=1 +x_0=500000 +y_0=0 +ellps=GRS80 +units=m +no_defs',
    4326: '+proj=longlat +datum=WGS84 +no_defs',
    4490: '+proj=longlat +ellps=GRS80 +no_defs',
    3857: '+proj=merc +a=6378137 +b=6378137 +lat_ts=0 +lon_0=0 +x_0=0 +y_0=0 +k=1 +units=m +no_defs',
}

LLM_MODEL = os.environ.get("CAISHA_LLM_MODEL", "qwen-plus")
VLM_MODEL = os.environ.get("CAISHA_VLM_MODEL", "qwen-vl-max")

BASIC_FIELDS = [
    ("location", "采砂区域位置（河流名、行政村、桩号范围等原文描述）"),
    ("enterprise", "开采企业名称"),
    ("mining_length", "可采区长度（含单位，如 1.2km）"),
    ("mining_area", "开采面积（含单位，如 15.6万m2）"),
    ("avg_width", "平均开采宽度（含单位）"),
    ("mining_period", "采砂期限（起止时间原文）"),
    ("licensed_volume", "许可开采总量（含单位，如 30万t）"),
    ("annual_control_volume", "年度采砂控制总量（含单位）"),
    ("ctrl_bottom_elevation", "开采控制底高程（含单位，如 22.41～22.19m）"),
    ("mining_method", "采砂作业方式"),
    ("mining_equipment", "采砂机具（种类与数量）"),
    ("stake_range", "采区桩号范围（如 K12+300～K13+500）"),
]

MISSING = "资料未见"


@register_tool('caisha_report_tool')
class CaishaReportTool(BaseTool):
    description = '''
    采砂场监测分析报告生成工具。输入砂场名称（和许可证号），自动完成：
    实施方案知识库检索抽取采砂场基本情况（11项）、现场监测报告整理、
    无人机正射影像裁剪+视觉模型解译（停工情况/作业器具/堆砂堆靠）、
    无人船水下测量高程与控制开采高程比对评估、高程点示意图绘制，
    最终生成 Word 监测报告并返回下载链接。
    前置条件：该砂场 RTK 测点已通过 scripts/import_rtk_dat.py 入库 rtk_points 表。
    '''

    parameters = [
        {'name': 'site_name', 'type': 'string',
         'description': '砂场名称（须与 rtk_points.site_name、2026年采区.shp 的 Name 一致，如：郝楼砂场）',
         'required': True},
        {'name': 'permit_no', 'type': 'string',
         'description': '许可证号（如：豫潢砂许〔2025〕第01号）；不传则从 rtk_points 表读取',
         'required': False},
        {'name': 'kb_area_query', 'type': 'string',
         'description': '实施方案检索用采区名（砂场名与实施方案采区名不一致时指定，如黄寨砂场→黄堰村采区）',
         'required': False},
        {'name': 'plan_year', 'type': 'string',
         'description': '实施方案年度（默认 2025）', 'required': False},
        {'name': 'tif_path', 'type': 'string',
         'description': f'无人机正射影像路径（默认 {DEFAULT_TIF}）', 'required': False},
        {'name': 'tif_epsg', 'type': 'integer',
         'description': '影像坐标系头信息损坏时的兜底 EPSG（默认 4547）', 'required': False},
    ]

    # ------------------------------------------------------------------
    # 站点类型 / 规范名 / 许可证号（标题、图注统一使用）
    # ------------------------------------------------------------------
    @staticmethod
    def _is_project(site: str) -> bool:
        """疏浚/修复/治理等工程类站点判定（非采砂场）。"""
        project_kw = ('工程', '项目', '修复', '疏浚', '清淤', '治理', '整治', '扩容')
        sand_kw = ('采区', '砂场', '可采', '采点')
        return any(k in site for k in project_kw) and not any(k in site for k in sand_kw)

    @staticmethod
    def _display_name(site: str, is_project: bool) -> str:
        """规范站点名：工程类去掉"南-南1 / -01"等分段后缀，全部去掉 RTK 冗余。"""
        name = site.replace('RTK', '').replace('rtk', '').strip()
        if is_project:
            prev = None
            while prev != name:
                prev = name
                name = re.sub(r'[东南西北中][-_][东南西北中0-9]+$', '', name).strip()
                name = re.sub(r'[-_][A-Za-z0-9]{1,4}$', '', name).strip()
        return name or site

    @staticmethod
    def _normalize_permit(permit_no: str) -> str:
        """许可证号规范化：各种括号统一为〔〕，去空格。"""
        if not permit_no or permit_no == MISSING:
            return MISSING
        p = re.sub(r'[\[(（]', '〔', permit_no)
        p = re.sub(r'[\])）]', '〕', p)
        p = re.sub(r'\s+', '', p)
        return p or MISSING

    # ------------------------------------------------------------------
    # 入口
    # ------------------------------------------------------------------
    def call(self, params: Union[str, Dict[str, Any]], **kwargs) -> Dict[str, Any]:
        if isinstance(params, str):
            try:
                params = json.loads(params)
            except json.JSONDecodeError:
                return {'success': False, 'error': 'Invalid JSON parameters'}

        site = (params.get('site_name') or '').strip()
        if not site:
            return {'success': False, 'error': 'site_name 不能为空'}
        # 站点名归一化：同许可证多采区（青龙闵塆采区-1/-2）应合并为一份报告，
        # 输入带 "-数字" 后缀别名时，若主名在 rtk_points 存在则归一化到主名
        site = self._normalize_site_name(site)
        plan_year = str(params.get('plan_year') or '2025')
        notes: List[str] = []
        tif_path = params.get('tif_path')
        tif_source = ''  # 记录影像来源便于排查
        if not tif_path:
            # 优先用砂场专属彩色影像（全县大图是灰度的）
            # 第一步：精确文件名匹配 {site}.tif
            site_tif = os.path.join(SITE_TIF_DIR, f'{site}.tif')
            if os.path.exists(site_tif):
                tif_path = site_tif
                tif_source = '精确匹配'
            else:
                # 第二步：子目录中搜索同名文件
                for root, dirs, files in os.walk(SITE_TIF_DIR):
                    for f in files:
                        if f == f'{site}.tif':
                            site_tif = os.path.join(root, f)
                            break
                    if os.path.exists(site_tif):
                        break
                if os.path.exists(site_tif):
                    tif_path = site_tif
                    tif_source = '子目录精确匹配'
                else:
                    # 第三步：模糊匹配——在子目录中找包含站点关键词的 .tif
                    # 提取站点核心词（去掉采区/可采区/部分河段/镇/村等后缀）
                    import re as _re
                    core_kw = _re.sub(r'(可采区|采区|可采点|部分河段|砂场)', '', site).strip()
                    # 分段：从核心词中提取两个以上连续汉字作为匹配关键词
                    kw_candidates = [site]
                    if core_kw != site:
                        kw_candidates.append(core_kw)
                    # 对含顿号的站点名（如 楠杆镇李寨村、邵湾村部分河段），拆分后也加入
                    for sep in ('、', '，', ',', '和'):
                        parts = site.split(sep)
                        if len(parts) >= 2:
                            kw_candidates.extend(p.strip() for p in parts if len(p.strip()) >= 3)

                    best_tif = ''
                    best_score = 0
                    for root, dirs, files in os.walk(SITE_TIF_DIR):
                        for f in files:
                            if not f.lower().endswith('.tif'):
                                continue
                            f_no_ext = _re.sub(r'\.tif$', '', f, flags=_re.IGNORECASE)
                            for kw in kw_candidates:
                                if kw in f_no_ext:
                                    score = len(kw) / max(len(f_no_ext), 1) + (1 if kw == f_no_ext else 0)
                                    if score > best_score:
                                        best_score = score
                                        best_tif = os.path.join(root, f)
                                    break
                    if best_tif and os.path.exists(best_tif):
                        tif_path = best_tif
                        tif_source = f'模糊匹配({os.path.basename(best_tif)})'
                    else:
                        # 第四步：在子目录名/文件中匹配站点关键词，取该目录下任意 .tif
                        for root, dirs, files in os.walk(SITE_TIF_DIR):
                            # 检查目录名或目录内任意文件名是否含站点关键词
                            dir_match = any(kw in os.path.basename(root) for kw in kw_candidates)
                            file_match = any(
                                kw in f for f in files for kw in kw_candidates
                                if not f.lower().endswith('.xml') and not f.lower().endswith('.db'))
                            if dir_match or file_match:
                                tifs_in_dir = [f for f in files
                                               if f.lower().endswith('.tif')
                                               and not f.lower().endswith('.tif.aux.xml')]
                                if tifs_in_dir:
                                    tif_path = os.path.join(root, tifs_in_dir[0])
                                    tif_source = f'目录关键词匹配({os.path.basename(tif_path)})'
                                    break
                                # 无 .tif 但目录包含站点关键词，继续找
                        if not tif_path:
                            tif_path = DEFAULT_TIF
                            tif_source = '默认影像(潢川县.tif)'
            if tif_source:
                notes.append(f'正射影像来源：{tif_source}')
        tif_epsg = int(params.get('tif_epsg') or 4547)
        area_query = (params.get('kb_area_query') or '').strip()

        # ---- 数据库测点统计（同时拿到许可证号/控制高程）----
        stats = self._rtk_stats(site)
        if not stats:
            return {'success': False,
                    'error': f'rtk_points 表中无砂场"{site}"的测点，请先运行 scripts/import_rtk_dat.py 入库'}
        permit_no = self._normalize_permit(params.get('permit_no') or stats['permit_no'] or MISSING)

        # ---- 站点类型与规范名（标题/图注统一使用）----
        is_project = self._is_project(site)
        display_name = self._display_name(site, is_project)
        site_title = display_name if is_project else (permit_no if permit_no != MISSING else display_name)

        # ---- 1. 基本情况（实施方案 + LLM 抽取/散文生成）----
        area_name = area_query or self._area_name_from_source(stats, site)
        if is_project and not area_query:
            area_name = display_name
        basic, basic_src = self._extract_basic_info(area_name, site)
        # 工程类项目（防洪/疏浚/修复等，非采砂场）：采砂实施方案的通用条款（开采期限、
        # 机具、总量等）不适用，强制置“资料未见”，防止 LLM 把采砂条款串扰进工程报告
        if is_project:
            for _f in ('mining_period', 'mining_method', 'mining_equipment',
                       'licensed_volume', 'annual_control_volume',
                       'mining_length', 'mining_area', 'avg_width'):
                basic[_f] = MISSING
        basic_prose, prose_src = self._extract_basic_prose(area_name, site, is_project, basic)
        if basic_src:
            notes.append(f'基本情况来源：{basic_src}')
        if prose_src and prose_src != basic_src:
            notes.append(f'基本情况散文来源：{prose_src}')

        # ---- 2. 现场监测（现场监测报告，含现场图片）----
        field_monitoring, field_photos = self._field_monitoring(site, area_name, display_name)

        # ---- 3. 影像裁剪 + VLM 解译 ----
        out_dir = os.path.join(BACKEND_DIR, 'static', 'reports')
        os.makedirs(out_dir, exist_ok=True)
        ts = int(time.time())
        ortho_png = os.path.join(out_dir, f'caisha_{ts}_ortho.png')
        clip = self._clip_drone_image(site, tif_path, tif_epsg, ortho_png,
                                      points_4326=stats.get('points'))
        # 砂场专属 TIF 裁剪失败时，回退默认影像（全县灰度大图）
        if clip is None and tif_path != DEFAULT_TIF and os.path.exists(DEFAULT_TIF):
            logger.info('砂场影像裁剪失败，尝试默认影像 %s', DEFAULT_TIF)
            clip = self._clip_drone_image(site, DEFAULT_TIF, tif_epsg, ortho_png,
                                          points_4326=stats.get('points'))
        drone_interp = MISSING
        if clip:
            drone_interp = self._vlm_interpret(ortho_png, site, permit_no)
        else:
            ortho_png = None
            notes.append('无人机影像裁剪失败（影像缺失/范围不符），报告中无正射影像图')

        # ---- 4. 深度评估 ----
        depth_evaluation, ctrl_text = self._depth_evaluation(stats)
        if ctrl_text:
            basic['ctrl_bottom_elevation'] = ctrl_text  # 实测入库的控制高程优先

        # ---- 5. 高程点示意图 ----
        elev_png = os.path.join(out_dir, f'caisha_{ts}_elev.png')
        if not self._draw_elev_figure(site, stats['points'], clip, elev_png):
            elev_png = None
            notes.append('高程点示意图绘制失败')

        # ---- 组装变量并填充模板 ----
        # 图号顺延：现场监测图片占前面的图号，正射/高程图图号动态计算
        n_field = len(field_photos)
        variables = dict(basic)
        variables.update({
            'site_name': site,
            'site_title': site_title,
            'display_name': display_name,
            'permit_no': permit_no,
            'plan_year': plan_year,
            'basic_info': basic_prose,
            'basic_section_name': '疏浚工程基本情况' if is_project else '采砂场基本情况',
            'monitor_section_name': '疏浚工程监测实施情况' if is_project else '采砂场监测实施情况',
            'drone_subject': display_name if is_project else f'{permit_no}采砂区域',
            'drone_scope': ('批复工程范围' if is_project
                            else f'{plan_year}年度规划批复范围（桩号{basic.get("stake_range", MISSING)}）'),
            'field_monitoring': field_monitoring,
            'drone_interpretation': drone_interp,
            'depth_evaluation': depth_evaluation,
            'point_count': str(stats['count']),
            'elev_min': f"{stats['zmin']:.3f}",
            'elev_max': f"{stats['zmax']:.3f}",
            'elev_avg': f"{stats['zavg']:.3f}",
            'fig_ortho_no': str(n_field + 1),
            'fig_elev_no': str(n_field + 2),
            'generated_date': time.strftime('%Y年%m月%d日'),
        })
        for key, _ in BASIC_FIELDS:
            variables.setdefault(key, MISSING)

        try:
            result = self._render_docx(variables, ortho_png, elev_png, field_photos)
        except Exception as e:
            logger.exception('caisha report render failed')
            return {'success': False, 'error': f'报告渲染失败: {e}'}
        result['notes'] = notes
        result['variables'] = {k: v for k, v in variables.items() if not k.startswith('_')}
        return result

    # ------------------------------------------------------------------
    # 数据库
    # ------------------------------------------------------------------
    def _normalize_site_name(self, site: str) -> str:
        """站点名归一化：输入在 rtk_points 中不存在、且形如 "主名-数字"（如 青龙闵塆采区-1）时，
        若去掉 "-数字" 后缀的主名存在测点，则归一化到主名（同许可证多采区合并为一份报告）。
        合法含 "-数字" 后缀的站点名（如 竹竿河-1、固始县...工程-01）优先按原样匹配。"""
        import psycopg2

        def exists(name: str) -> bool:
            try:
                conn = psycopg2.connect(**DB_CFG)
                cur = conn.cursor()
                cur.execute("SELECT count(*) FROM rtk_points WHERE site_name = %s", (name,))
                n = cur.fetchone()[0]
                conn.close()
                return n > 0
            except Exception:
                return False

        if exists(site):
            return site
        m = re.match(r'^(.*?)[-_](\d+)$', site)
        if m and exists(m.group(1)):
            return m.group(1)
        return site

    def _rtk_stats(self, site: str) -> Optional[Dict[str, Any]]:
        import psycopg2
        conn = psycopg2.connect(**DB_CFG)
        cur = conn.cursor()
        cur.execute(
            """SELECT count(*), min(elev), avg(elev), max(elev),
                      max(permit_no), max(ctrl_elev_max), max(ctrl_elev_min), max(ctrl_elev_source)
               FROM rtk_points WHERE site_name = %s""", (site,))
        row = cur.fetchone()
        if not row or not row[0]:
            conn.close()
            return None
        cur.execute(
            """SELECT ST_X(geom), ST_Y(geom), elev FROM rtk_points
               WHERE site_name = %s ORDER BY id""", (site,))
        points = [(float(x), float(y), float(z)) for x, y, z in cur.fetchall()]
        conn.close()
        return {
            'count': int(row[0]),
            'zmin': float(row[1]), 'zavg': float(row[2]), 'zmax': float(row[3]),
            'permit_no': row[4] or '',
            'ctrl_max': float(row[5]) if row[5] is not None else None,
            'ctrl_min': float(row[6]) if row[6] is not None else None,
            'ctrl_src': row[7] or '',
            'points': points,
        }

    @staticmethod
    def _area_name_from_source(stats: Dict[str, Any], site: str) -> str:
        """优先从入库时记录的控制高程来源标题里解析采区名（如 1.6.2 黄堰村采区）。"""
        m = re.search(r'([\u4e00-\u9fa5]{2,8}村?采区)', stats.get('ctrl_src') or '')
        if m:
            return m.group(1)
        base = site.replace('砂场', '')
        if base.endswith('采区'):
            return base
        return base + '采区'

    # ------------------------------------------------------------------
    # 知识库检索 + LLM 抽取
    # ------------------------------------------------------------------
    @staticmethod
    def _kb_search(query: str, project: str, top_k: int = 20) -> List[Dict[str, str]]:
        try:
            from tools.llamaindex_knowledge_tool import KnowledgeBaseTool
            kb = KnowledgeBaseTool()
            res = kb.call({'operation': 'search', 'query': query,
                           'top_k': top_k, 'project': project})
            if res.get('success'):
                return res.get('data') or []
        except Exception as e:
            logger.warning('KB search failed (%s/%s): %s', project, query, e)
        return []

    @staticmethod
    def _llm(prompt: str) -> str:
        import dashscope
        resp = dashscope.Generation.call(
            model=LLM_MODEL,
            api_key=os.environ.get('DASHSCOPE_API_KEY'),
            messages=[{'role': 'user', 'content': prompt}],
            result_format='message',
        )
        if resp.status_code == 200:
            return resp.output.choices[0].message.content.strip()
        raise RuntimeError(f'LLM 调用失败: {resp.message}')

    def _retrieve_basic_chunks(self, area_name: str) -> Tuple[List[str], List[str]]:
        """实施方案知识库检索：返回 (texts, srcs)。按候选名最长命中排序，
        只保留标题/正文中确实出现目标采区名的 chunk。"""
        base = area_name.replace('可采区', '').replace('采区', '')
        # 候选名（去分隔符/村后缀/河流前缀/行政区划前缀），用于排序时正确识别相关 chunk
        candidates_rank = {base, base.replace('、', '').replace('，', '').replace(',', '').replace('-', '')}
        # 按分隔符拆分为单独词（如 牛老家、龙潭 → 牛老家, 龙潭）
        for sep in ('、', '，', ',', '-'):
            for p in base.split(sep):
                p = p.strip()
                if len(p) >= 2:
                    candidates_rank.add(p)
        # 去掉通用后缀（部分河段/部分/河段），如 高店镇王湾村部分河段 → 高店镇王湾村
        for c in list(candidates_rank):
            for sfx in ('部分河段', '部分', '河段'):
                if c.endswith(sfx) and len(c) > len(sfx) + 1:
                    candidates_rank.add(c[:-len(sfx)])
        for c in list(candidates_rank):
            for sfx in ('村', '镇', '乡'):
                if c.endswith(sfx) and len(c) > len(sfx) + 1:
                    candidates_rank.add(c[:-len(sfx)])
        for c in list(candidates_rank):
            for r in ('史灌河', '史河', '灌河', '淮河', '潢河', '竹竿河', '白露河', '浉河', '迎河'):
                if c.startswith(r) and len(c) > len(r) + 1:
                    candidates_rank.add(c[len(r):])
                    break
        # 行政区划前缀剥离：高店镇王湾村 → 王湾村/王湾（匹配罗山 KB 摘要 "高店乡(王湾村)"）
        for c in list(candidates_rank):
            m = re.match(r'^[\u4e00-\u9fa5]{1,3}[镇乡](.+)$', c)
            if m and len(m.group(1)) >= 2:
                core = m.group(1)
                candidates_rank.add(core)
                candidates_rank.add(core.rstrip('村'))
        # 镇/乡互换：高店镇王湾村 ↔ 高店乡王湾村（罗山 MD 用乡，站点名用镇）
        for c in list(candidates_rank):
            for a, b in (('镇', '乡'), ('乡', '镇')):
                if a in c:
                    candidates_rank.add(c.replace(a, b))
        hits = self._kb_search(f'{area_name} 长度 面积 平均宽度 开采总量 采砂期限 开采企业 桩号 控制开采高程',
                               'shishifangan')
        hits += self._kb_search(f'{area_name} 采砂期限 开采企业 作业方式 采砂机具', 'shishifangan', top_k=10)
        # 县级通用条款（开采期限等）单独检索、强制附加，避免被排序挤掉
        extra = self._kb_search('开采期限 采砂期限', 'shishifangan', top_k=3)
        # 优先含采区名的 chunk（按候选名最长命中排序，稳定保留 KB 相关度顺序），去重后最多取 10 段
        seen = set()
        uniq = []
        for h in hits:
            key = (h.get('title', ''), (h.get('content', '') or '')[:80])
            if key not in seen:
                seen.add(key)
                uniq.append(h)

        def _match_len(h: Dict[str, str]) -> int:
            text = h.get('title', '') + h.get('content', '')
            return max((len(c) for c in candidates_rank if c and c in text), default=0)

        ranked = sorted(uniq, key=lambda h: -_match_len(h))
        # 只保留标题/正文中确实出现目标采区名的 chunk，避免其他采区（如相邻 LX1）内容被 LLM 当作目标抽取
        picked = [h for h in ranked[:10] if _match_len(h) > 0]
        picked_keys = {(h.get('title', ''), (h.get('content', '') or '')[:80]) for h in picked}
        for h in extra[:2]:
            key = (h.get('title', ''), (h.get('content', '') or '')[:80])
            if key not in picked_keys:
                picked_keys.add(key)
                picked.append(h)
        texts, srcs = [], []
        for h in picked:
            texts.append(f"【{h.get('title', '')}】\n{h.get('content', '')}")
            srcs.append(h.get('title', ''))
        return texts, srcs

    def _extract_basic_info(self, area_name: str, site: str) -> Tuple[Dict[str, str], str]:
        """结构化抽取 11+1 个变量（散文生成失败时的兜底，也是深度评估字段的来源）。"""
        texts, srcs = self._retrieve_basic_chunks(area_name)
        if not texts:
            logger.warning('实施方案未检索到采区 %s 相关内容，尝试本地MD表格', area_name)
            # KB搜索无结果 → 兜底：本地MD表格
            return self._ssfa_lookup_result(site), ''

        field_desc = '\n'.join(f'- {k}: {d}' for k, d in BASIC_FIELDS)
        prompt = f"""你是水利行业资料整理员。以下是河道采砂实施方案中与"{area_name}"（对应{site}）有关的原文片段。
请只依据原文抽取该采区的信息，严格保留原文数值与单位，不得推算、不得编造；原文没有的字段填"{MISSING}"。

需要抽取的字段：
{field_desc}

注意：采砂期限、作业时间等县级通用条款（如"3.2.1 开采期限"）同样适用于本采区，可作为对应字段的依据。
但严禁把其他采区（如LX1淮河村可采区、LG9吴乡村可采区等）的长度、宽度、高程、开采量等信息当作本采区的信息；若没有任何片段明确描述本采区/工程，则所有字段一律填"{MISSING}"。

原文片段：
{chr(10).join(texts)}

只输出一个 JSON 对象（不要 Markdown 代码块），key 为上述英文字段名，value 为中文字符串。"""
        try:
            raw = self._llm(prompt)
            raw = re.sub(r'^```(json)?|```$', '', raw.strip(), flags=re.MULTILINE).strip()
            data = json.loads(raw)
            basic = {k: str(data.get(k) or MISSING).strip() for k, _ in BASIC_FIELDS}
            # 如果LLM抽取结果大部分是"资料未见"，尝试本地MD表格补充
            missing_count = sum(1 for v in basic.values() if v == MISSING)
            if missing_count >= 6:
                logger.info('LLM抽取缺项过多(%d/11)，尝试本地MD表格补充', missing_count)
                md_data = self._ssfa_lookup_result(site)
                if md_data:
                    # 只补充KB没有的字段（KB结果优先）
                    for k, v in md_data.items():
                        if v and basic.get(k) == MISSING:
                            basic[k] = v
            return basic, '; '.join(dict.fromkeys(srcs))
        except Exception as e:
            logger.warning('基本情况抽取失败: %s', e)
            # LLM失败也尝试MD表格兜底
            md_fallback = self._ssfa_lookup_result(site)
            return (md_fallback or {}), '; '.join(dict.fromkeys(srcs))

    def _extract_basic_prose(self, area_name: str, site: str, is_project: bool,
                             basic: Dict[str, str]) -> Tuple[str, str]:
        """按《2025采砂报告》参考文风生成"基本情况"散文段：只依据知识库原文，
        严格保留数值单位，缺失写"资料未见"，严禁编造。失败/无检索结果时退回兜底句式。"""
        display_name = self._display_name(site, is_project)
        texts, srcs = self._retrieve_basic_chunks(area_name)
        if not texts:
            logger.warning('实施方案未检索到 %s，基本情况散文退化为兜底句式', area_name)
            return self._build_basic_fallback(site, is_project, basic), ''
        if is_project:
            example = ('潢川县白露河上游防洪能力提升工程位于白露河桩号77+600（沪陕高速桥）～102+600'
                       '（潢川光山县界）之间，本次砂石处置只对白露河干流（桩号：80+685～82+375 段、'
                       '83+359～85+690 段）潢川段河道清淤疏浚工程中产生的48.3万m3弃土（渣）进行处置，'
                       '处置单位为潢川城投建材有限公司，处置周期10个月。河道清淤疏浚弃土采用挖掘机开挖，'
                       '设计疏浚后控制高程位于41.13m-52.65m之间。')
            rules = ('本工程为疏浚/修复/治理类工程，应整理工程概况要素：工程位置、桩号范围、'
                     '清淤疏浚内容、弃土（渣）处置量、处置单位、处置周期、开挖方式、设计控制高程等；'
                     '严禁套用实施方案中其他采区的采砂条款（采砂期限、采砂机具、许可开采总量等一律不写）。')
        else:
            example = ('郝楼砂场采砂区域位于潢河魏岗镇郝楼村河段，属规划潢河郝楼村可采区HHK1#范围；'
                       '左岸-魏岗镇郝楼村；右岸-谈店乡老君台村，小吕河村，开采企业为河南水投光州现代水网'
                       '有限公司。可采区长度821m，开采面积17.04万㎡。采砂期限为2025年发证后至2026年6月14日，'
                       '许可开采总量20.01万m³，年度采砂控制总量20.01万m³，开采控制底高程22.41～22.19m。'
                       '采砂作业方式为旱采与水采相结合，采砂机具为2艘采砂船、3艘运砂船、1辆挖掘机、'
                       '2辆铲车及2套砂石分离系统。')
            rules = ('应整理要素：采砂区域位置（河流/村/河段）、规划采区编号、左右岸、开采企业、'
                     '可采区长度、开采面积、平均开采宽度、采砂期限、许可开采总量、年度采砂控制总量、'
                     '开采控制底高程、采砂作业方式、采砂机具；采砂期限等县级通用条款适用于本采区时可采用。')
        prompt = f"""你是水利行业资料整理员。以下是河道采砂实施方案中与"{area_name}"（对应{site}）有关的原文片段。
请参照下面示例的句式与详略程度，整理成"{("疏浚工程基本情况" if is_project else "采砂场基本情况")}"一节的正文段落（1段，不加标题、不加编号）：

示例：
{example}

要求：
1. 只依据原文事实，严格保留原文数值与单位，不得推算、不得编造；
2. {rules}
3. 原文没有的信息一律写"{MISSING}"，不杜撰；
4. 严禁把其他采区（如LX1淮河村可采区、LG9吴乡村可采区等）的长度、宽度、高程、开采量等信息写入本段；
5. 若原文没有任何与本工程/采区直接相关的内容，整段只写一句"……（{MISSING}）"。

原文片段：
{chr(10).join(texts)}

只输出正文段落本身（不要 Markdown、不要 JSON、不要任何标题或解释）。"""
        try:
            prose = self._llm(prompt).strip()
            prose = prose.replace(chr(10), '').strip()
            if len(prose) < 15 or prose.startswith('{') or prose.startswith(chr(96) * 3):
                raise ValueError(f'散文生成结果异常: {prose[:40]}')
            return prose, '; '.join(dict.fromkeys(srcs))
        except Exception as e:
            logger.warning('基本情况散文生成失败: %s，退化为兜底句式', e)
            return self._build_basic_fallback(site, is_project, basic), ''

    def _build_basic_fallback(self, site: str, is_project: bool,
                              basic: Dict[str, str]) -> str:
        """散文生成失败/无检索结果时的兜底段落（旧版固定句式）。"""
        display_name = self._display_name(site, is_project)
        if is_project:
            if basic.get('location') and basic['location'] != MISSING:
                return f'{display_name}位于{basic["location"]}。'
            return f'{display_name}实施方案知识库中未检索到该工程概况（{MISSING}）。'
        parts = [
            f'{display_name}采砂区域位于{basic.get("location", MISSING)}，开采企业为{basic.get("enterprise", MISSING)}。',
            f'可采区长度{basic.get("mining_length", MISSING)}，开采面积{basic.get("mining_area", MISSING)}，'
            f'平均开采宽度{basic.get("avg_width", MISSING)}。',
            f'采砂期限为{basic.get("mining_period", MISSING)}，许可开采总量{basic.get("licensed_volume", MISSING)}，'
            f'年度采砂控制总量{basic.get("annual_control_volume", MISSING)}，'
            f'开采控制底高程{basic.get("ctrl_bottom_elevation", MISSING)}。',
            f'采砂作业方式为{basic.get("mining_method", MISSING)}，采砂机具为{basic.get("mining_equipment", MISSING)}。',
        ]
        return ''.join(parts)

    def _ssfa_lookup_result(self, site: str) -> Dict[str, str]:
        """从本地MD表格提取数据并转换为报告字段格式。

        返回的键名与 BASIC_FIELDS 对齐（mining_length / mining_area /
        ctrl_bottom_elevation / annual_control_volume / mining_equipment 等），
        可直接合并到 basic 字典中。
        """
        row = self._ssfa_lookup(site)
        if not row:
            return {}
        result = {}
        # 可采区长度
        if row.get('len'):
            result['mining_length'] = f"{row['len']}m"
        # 开采面积
        if row.get('area'):
            result['mining_area'] = f"{row['area']}km²"
        # 开采控制底高程
        if row.get('ctrl_elev'):
            result['ctrl_bottom_elevation'] = f"{row['ctrl_elev']}m"
        # 年度采砂控制总量
        if row.get('ctrl_qty'):
            result['annual_control_volume'] = f"{row['ctrl_qty']}万m³"
        # 采砂机具：格式化为文字描述
        parts = []
        if row.get('ship'):
            parts.append(f"采砂船{row['ship']}艘")
        if row.get('lift'):
            parts.append(f"提砂船{row['lift']}艘")
        if row.get('digger'):
            parts.append(f"挖掘机{row['digger']}辆")
        if parts:
            result['mining_equipment'] = '，'.join(parts)
        # 采砂区域位置（行政位置）
        if row.get('location'):
            result['location'] = row['location']
        # 采区桩号范围
        if row.get('stake_range'):
            result['stake_range'] = row['stake_range']
        # 储砂场 → 可作为开采企业的参考
        if row.get('storage'):
            result['enterprise'] = row['storage']
        return result

    # ------------------------------------------------------------------
    # 本地 MD 表格解析（兜底：KB embedding 对 HTML 表格检索效果差）
    # ------------------------------------------------------------------
    _GUSHI_SSFA_MD = os.path.join(os.path.dirname(os.path.dirname(BACKEND_DIR)), '采砂', '固始县实施方案.md')
    _SSFA_TABLE_CACHE = None  # 缓存解析结果

    @classmethod
    def _parse_ssfa_table(cls) -> List[Dict[str, str]]:
        """解析固始县实施方案.md 中的综合表格，缓存结果。

        表格存在 colspan/rowspan 导致各行列数不一致的复杂情况，
        因此改用内容模式匹配代替固定列偏移，更鲁棒。
        """
        if cls._SSFA_TABLE_CACHE is not None:
            return cls._SSFA_TABLE_CACHE
        if not os.path.exists(cls._GUSHI_SSFA_MD):
            logger.warning('实施方案MD文件不存在: %s', cls._GUSHI_SSFA_MD)
            cls._SSFA_TABLE_CACHE = []
            return []
        with open(cls._GUSHI_SSFA_MD, encoding='utf-8') as f:
            content = f.read()
        # 找到综合表格：包含 "开采点名称" 的单个 <table> 块（避免跨表匹配）
        blocks = content.split('</table>')
        table_html = None
        for blk in blocks:
            if '开采点名称' in blk and '年度采砂控制量' in blk:
                idx = blk.rfind('<table')
                if idx >= 0:
                    table_html = blk[idx:] + '</table>'
                    break
        if not table_html:
            logger.warning('实施方案MD中未找到综合表格')
            cls._SSFA_TABLE_CACHE = []
            return []
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL)
        if len(rows) < 3:
            cls._SSFA_TABLE_CACHE = []
            return []

        results = []
        for row in rows[2:]:  # skip 2 header rows
            cells = [c.strip() for c in re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', row, re.DOTALL)]
            if len(cells) < 4:
                continue
            name = cells[1] if len(cells) > 1 else ''
            if not name or name in ('合计', '小计', ''):
                continue

            entry = {'_row_name': name}

            # --- 内容模式匹配：按单元格内容特征定位各字段 ---

            # 河段长度：3-4位纯数字（600~4000），且不在桩号/高程范围内
            len_idx = -1
            for i, c in enumerate(cells):
                if re.match(r'^\d{3,4}$', c):
                    n = int(c)
                    if 500 <= n <= 4000:
                        # 确保后面的单元格像是面积（小数）
                        if i + 1 < len(cells) and re.match(r'^\d+\.\d+$', cells[i + 1]):
                            len_idx = i
                            break
            if len_idx >= 0:
                entry['len'] = cells[len_idx]
                # 河段面积 = 紧接在长度后面的小数
                if len_idx + 1 < len(cells) and re.match(r'^\d+\.\d+$', cells[len_idx + 1]):
                    entry['area'] = cells[len_idx + 1]

            # 年度控制开采高程：必须出现在河段长度之后（前面的 XX.XX-XX.XX 是深泓线/规划高程）
            ctrl_elev_idx = -1
            for i, c in enumerate(cells):
                if (len_idx < 0 or i > len_idx) and re.match(r'^\d+\.\d+\s*[-~]\s*\d+\.\d+$', c):
                    entry['ctrl_elev'] = c
                    ctrl_elev_idx = i
                    # 年度采砂控制量：控制高程之后最近一个 >1 的小数（排除可利用系数 0.x）
                    for j in range(i + 1, min(i + 6, len(cells))):
                        if re.match(r'^\d+\.\d+$', cells[j]):
                            try:
                                if float(cells[j]) > 1.0:
                                    entry['ctrl_qty'] = cells[j]
                                    break
                            except ValueError:
                                pass
                    break

            # 储砂场：以"储砂场"结尾
            for c in cells:
                if c.endswith('储砂场'):
                    entry['storage'] = c
                    break

            # 行政位置：靠前的中文长文本（含 村/镇/乡/社区 等，不含数字和+号）
            for i, c in enumerate(cells[2:9]):
                if len(c) >= 4 and ('村' in c or '镇' in c or '乡' in c or '社区' in c or '办事处' in c):
                    if not re.search(r'[\d+]', c):
                        entry['location'] = c
                        break

            # 2025计划采砂范围：取最靠近河段长度的桩号（前面的桩号是总规划范围）
            best_stake, best_dist = None, 999
            for i, c in enumerate(cells):
                if re.match(r'^\d+\+[\d]+\s*[-~]\s*\d+\+[\d]+$', c):
                    if len_idx >= 0:
                        dist = abs(i - len_idx)
                        if dist < best_dist:
                            best_dist, best_stake = dist, c
                    elif best_stake is None:
                        best_stake = c
            if best_stake:
                entry['stake_range'] = best_stake

            # 采砂船/提砂船/挖掘机：末尾连续的 1-2 位整数
            nums = []
            for c in reversed(cells):
                if re.match(r'^\d{1,2}$', c):
                    nums.append(c)
                else:
                    break
            nums.reverse()
            if len(nums) >= 3:
                entry['ship'], entry['lift'], entry['digger'] = nums[-3], nums[-2], nums[-1]
            elif len(nums) == 2:
                entry['lift'], entry['digger'] = nums[-2], nums[-1]
            elif len(nums) == 1:
                entry['digger'] = nums[-1]

            if entry.get('len') or entry.get('ctrl_elev'):
                results.append(entry)

        cls._SSFA_TABLE_CACHE = results
        logger.info('解析实施方案MD表格：%d 行数据', len(results))
        return results

    def _ssfa_lookup(self, site: str) -> Optional[Dict[str, str]]:
        """从本地MD表格中匹配砂场，返回提取的数据字段。"""
        rows = self._parse_ssfa_table()
        if not rows:
            return None
        # 生成站点候选名（复用 _field_monitoring 的逻辑）
        core = site.replace('砂场', '').replace('RTK', '').replace('可采区', '').replace('采区', '')
        candidates = {core}
        candidates.add(core.replace('、', '').replace('，', '').replace(',', '').replace('-', ''))
        for sep in ('、', '，', ',', '-'):
            for p in core.split(sep):
                p = p.strip()
                if len(p) >= 2:
                    candidates.add(p)
        for c in list(candidates):
            for sfx in ('村', '镇', '乡'):
                if c.endswith(sfx) and len(c) > len(sfx) + 1:
                    candidates.add(c[:-len(sfx)])
        for c in list(candidates):
            for r in ('史灌河', '史河', '灌河', '淮河', '潢河'):
                if c.startswith(r) and len(c) > len(r) + 1:
                    candidates.add(c[len(r):])
                    break
        best, best_score = None, 0
        for row in rows:
            rn = row.get('_row_name', '')
            score = sum(1 for c in candidates if c and c in rn)
            if score > best_score:
                best_score, best = score, row
        if best and best_score > 0:
            logger.info('SSFA表格匹配: %s → %s (score=%d)', site, best.get('_row_name'), best_score)
        return best

    def _field_monitoring(self, site: str, area_name: str,
                          display_name: Optional[str] = None) -> Tuple[str, List[str]]:
        """返回（现场监测正文, 现场图片本地路径列表）。图片取自命中切片的 image_paths 元数据。"""
        # 生成候选匹配名：KB 标题常省略 村/河流前缀/可采区 等，需生成多个变体
        # 注意：必须先替换 可采区 再替换 采区，否则 可采区→可（残留）
        core = site.replace('砂场', '').replace('RTK', '').replace('可采区', '').replace('采区', '')
        candidates = {core}
        # 去分隔符
        candidates.add(core.replace('、', '').replace('，', '').replace(',', '').replace('-', ''))
        # 按分隔符拆分独立词
        for sep in ('、', '，', ',', '-'):
            for p in core.split(sep):
                p = p.strip()
                if len(p) >= 2:
                    candidates.add(p)
        area_clean = area_name.replace('可采区', '').replace('采区', '')
        candidates.add(area_clean)
        candidates.add(area_clean.replace('、', '').replace('，', '').replace(',', '').replace('-', ''))
        # 去 村/镇/乡 后缀（KB 常省略，如 王楼村→王楼）
        for c in list(candidates):
            for sfx in ('村', '镇', '乡'):
                if c.endswith(sfx) and len(c) > len(sfx) + 1:
                    candidates.add(c[:-len(sfx)])
        # 去河流前缀（KB 标题常省略，如 史河汪营→汪营）
        _RIVERS = ['史灌河', '史河', '灌河', '淮河', '潢河', '竹竿河', '白露河', '浉河', '迎河', '灌河']
        for c in list(candidates):
            for r in _RIVERS:
                if c.startswith(r) and len(c) > len(r) + 1:
                    candidates.add(c[len(r):])
                    break
        # 排除纯河流/地理泛名（如 淮河村→淮河，会误匹配任何淮河*条目）
        _BLACKLIST = {'淮河', '灌河', '史河', '潢河', '浉河', '竹竿河', '白露河', '迎河', '史灌河', '河流', '河段', '部分'}
        candidates = {c for c in candidates if c not in _BLACKLIST}
        # 括号型站名（如 淮河(竹竿河入淮口至桃花岛段)航道疏浚及岸线整治工程）：
        # KB 标题常改写括号内文字（入淮口→入河口、省略段/及岸线整治等），需按去括号全名
        # 做 ≥3 字子串重叠匹配。通用工程词（工程/提升/整治/航道等）子串不计入，
        # 否则“淮河航道提升工程”等不相关章节标题会因“提升工/升工程”子串被误判相关
        core_nobracket = re.sub(r'[()（）]', '', core)
        subs_all = {core_nobracket[i:i + l] for l in range(3, len(core_nobracket) + 1)
                    for i in range(len(core_nobracket) - l + 1)}
        _GENERIC = ('工程', '河道', '提升', '整治', '项目', '建设', '施工', '管理', '范围',
                    '疏浚', '岸线', '航道', '修复', '防洪', '能力', '河段', '现场', '监测',
                    '采砂', '监管')
        subs = {s for s in subs_all if not any(w in s for w in _GENERIC)}  # 特征子串
        subs4 = {s for s in subs if len(s) >= 4}

        def _name_overlap(text: str) -> bool:
            hit = sum(1 for s in subs if s in text)
            hit4 = sum(1 for s in subs4 if s in text)
            return hit >= 2 or hit4 >= 1

        hits = self._kb_search(f'{area_name} 现场监测 巡查 检查情况', 'xianchangjiance')
        def _any_candidate_in(text: str) -> bool:
            return any(c and c in text for c in candidates)
        related = [h for h in hits
                   if _any_candidate_in(h.get('title', '') + h.get('content', ''))
                   or _name_overlap(h.get('title', '') + h.get('content', ''))]
        if not related:
            logger.warning('现场监测报告未检索到 %s 相关内容', site)
            return f'现场监测报告知识库中未检索到{site}的监测记录（{MISSING}）。', []
        # 切片元数据里的现场图片地址（逗号分隔）→ 下载到本地
        photo_urls = []
        for h in related[:5]:
            raw = (h.get('metadata') or {}).get('image_paths') or ''
            for u in raw.split(','):
                u = u.strip()
                if u and u not in photo_urls:
                    photo_urls.append(u)
        photos = self._download_photos(photo_urls[:4])
        texts = [f"【{h.get('title', '')}】\n{h.get('content', '')}" for h in related[:5]]
        dn = display_name or site
        prompt = f"""以下是现场监测报告中与“{dn}”有关的原文片段。请整理为报告“现场监测”一节的正文（1段），
只使用原文事实（监测时间、监测内容、发现的问题、处理情况等），不得编造；用正式书面语，不要加标题。
首句建议采用“20XX年X月对{dn}开展现场监测，发现……”的句式（监测时间以原文为准）。

原文片段：
{chr(10).join(texts)}"""
        try:
            return self._llm(prompt), photos
        except Exception as e:
            logger.warning('现场监测整理失败: %s', e)
            return texts[0][:500], photos

    @staticmethod
    def _download_photos(urls: List[str]) -> List[str]:
        """下载现场监测图片到报告目录，返回成功的本地路径列表。"""
        import requests
        out_dir = os.path.join(BACKEND_DIR, 'static', 'reports')
        os.makedirs(out_dir, exist_ok=True)
        paths = []
        for i, url in enumerate(urls):
            try:
                r = requests.get(url, timeout=20)
                if r.status_code == 200 and r.content:
                    ext = os.path.splitext(url.split('?')[0])[1] or '.png'
                    p = os.path.join(out_dir, f'caisha_field_{int(time.time())}_{i}{ext}')
                    with open(p, 'wb') as f:
                        f.write(r.content)
                    paths.append(p)
                else:
                    logger.warning('现场图片下载失败 %s: HTTP %s', url, r.status_code)
            except Exception as e:
                logger.warning('现场图片下载失败 %s: %s', url, e)
        return paths

    # ------------------------------------------------------------------
    # 影像裁剪 / VLM 解译 / 出图
    # ------------------------------------------------------------------
    def _load_boundary(self, site: str):
        """返回 (EPSG:4326 的采区 GeoDataFrame) 或 None"""
        import re, geopandas as gpd
        if not os.path.exists(CAIQU_SHP):
            return None
        gdf = gpd.read_file(CAIQU_SHP)
        base = site.replace('砂场', '').replace('RTK', '')
        # 候选名称列表：原文 → 去后缀 → 去掉','前的乡镇部分
        candidates = [base]
        base_no_num = re.sub(r'[-_]\d+$', '', base)
        if base_no_num != base:
            candidates.append(base_no_num)
        # 提取纯采区名（去掉乡镇前缀，如 楠杆镇李寨村、邵湾村部分河段 → 李寨村、邵湾村部分河段）
        short = re.sub(r'^[^、,，]+[镇乡村](?:部分)?', '', base)
        if short and short != base:
            candidates.append(short)
        # shapefile 可能用 可采点/采区/可采区 不同后缀，统一用核心词匹配
        core = base.replace('可采区', '').replace('采区', '').replace('可采点', '').rstrip('-')

        sel = gdf[gdf['Name'].astype(str).str.contains(base, na=False)]
        if sel.empty:
            for cand in candidates:
                if cand == base:
                    continue
                sel = gdf[gdf['Name'].astype(str).str.contains(cand, na=False)]
                if not sel.empty:
                    logger.info('边界名称匹配：%s → %s', site, cand)
                    break
        # 反向匹配：site 名称包含 shapefile 中的名称
        if sel.empty and core:
            for _, row in gdf.iterrows():
                shp_name = str(row['Name'])
                shp_core = shp_name.replace('可采区', '').replace('采区', '').replace('可采点', '').rstrip('-')
                if core and shp_core and (core in shp_core or shp_core in core):
                    sel = gdf[gdf['Name'] == shp_name]
                    logger.info('边界名称反向匹配：%s ↔ %s', site, shp_name)
                    break
        if sel.empty:
            logger.warning('边界名称匹配失败，site=%s', site)
            # 大营-范营 尝试用相邻的 童庙 采区边界兜底
            if '大营' in site or '范营' in site:
                sel = gdf[gdf['Name'].str.contains('童庙', na=False)]
                if not sel.empty:
                    logger.info('大营-范营边界兜底：使用童庙采区边界')
            if sel.empty:
                return None
        if sel.crs is None:
            sel = sel.set_crs(4326)
        return sel.to_crs(4326)

    @staticmethod
    def _load_hedao(boundary):
        """河道管理范围线（hx.geojson 线要素，左岸/右岸分开，共469条）：
        全量读取、不裁切、不融合、不处理，原样画线（与前端地图"红线底图"一致）。
        返回 GeoSeries(4326) 或 None。"""
        import geopandas as gpd
        if not os.path.exists(HEDAO_GEOJSON):
            return None
        try:
            g = gpd.read_file(HEDAO_GEOJSON)
            if g.crs is None:
                g = g.set_crs(4326)
            else:
                g = g.to_crs(4326)
            return g.geometry
        except Exception as e:
            logger.warning('河道管理范围线读取失败: %s', e)
            return None

    @staticmethod
    def _crs_from_epsg_safe(epsg: int):
        """安全的 EPSG→CRS 转换。
        CGCS2000 3度带（4547/4548/4549）优先用显式"去带号"Proj4 参数，不查库：
        - 环境 PROJ 库版本混乱（share/proj 为旧 layout 1.2 被 GDAL 拒绝，
          pyproj 自带库对 4547 的定义与官方不一致），from_epsg 结果不确定；
        - 无人机影像与 RTK 数据坐标均为去带号米制（x≈550xxx），必须 x_0=500000
          才能匹配；标准带号假东（38 带=38500000）会错开一个带号、永远无交集。"""
        try:
            from rasterio.crs import CRS as RCRS
            proj4 = FALLBACK_PROJ4.get(epsg)
            if proj4:
                return RCRS.from_proj4(proj4)
            crs = RCRS.from_epsg(epsg)
            if crs is not None:
                return crs
        except Exception:
            pass
        from rasterio.errors import CRSError
        raise CRSError(f'Cannot resolve EPSG:{epsg} (PROJ db unavailable, no fallback)')

    def _clip_drone_image(self, site: str, tif_path: str, tif_epsg: int,
                          out_png: str,
                          points_4326: Optional[List[Tuple[float, float, float]]] = None
                          ) -> Optional[Dict[str, Any]]:
        """按采区范围（外扩30%）裁剪影像，叠加红色批复边界出 PNG。
        返回 {'crs':..., 'extent': (minx,maxx,miny,maxy), 'boundary_xy': [...]} 供高程图复用。
        points_4326：该站点 RTK 测点（4326）。测点包围盒超出批复边界时（多片测区，
        如 楠杆镇李寨村、邵湾村=SH4+SH5 三片，shp 仅有其中一片），并入裁剪窗口，
        保证报告影像/高程图覆盖全部实测区域。"""
        try:
            import numpy as np
            import rasterio
            from rasterio.crs import CRS
            from rasterio.enums import Resampling
            from rasterio.errors import WindowError
            from rasterio.windows import from_bounds
            from pyproj import Transformer

            boundary = self._load_boundary(site)
            if boundary is None and points_4326:
                # 无批复边界（工程类项目，如 白露河上游防洪能力提升工程）时，
                # 用 RTK 测点包围盒（外扩约110m）兜底裁剪窗口，保证正射图覆盖实测区域
                import geopandas as _gpd
                from shapely.geometry import box as _box
                minx = min(p[0] for p in points_4326); maxx = max(p[0] for p in points_4326)
                miny = min(p[1] for p in points_4326); maxy = max(p[1] for p in points_4326)
                boundary = _gpd.GeoDataFrame(
                    geometry=[_box(minx - 0.001, miny - 0.001, maxx + 0.001, maxy + 0.001)],
                    crs='EPSG:4326')
                logger.info('无批复边界，用测点包围盒兜底裁剪（%s）', site)
            if boundary is None:
                logger.warning('2026年采区.shp 中未找到 %s', site)
                return None
            if not os.path.exists(tif_path):
                logger.warning('影像不存在: %s', tif_path)
                return None

            # OVERVIEW_LEVEL=NONE：潢川县.tif 内部金字塔损坏（读出全零），
            # 必须绕开概览层、从基底原生数据降采样
            with rasterio.open(tif_path, OVERVIEW_LEVEL='NONE') as src:
                crs = src.crs
                # 头信息损坏（标经纬度但坐标是米制）→ 使用兜底 EPSG
                if crs is None or (crs.is_geographic and abs(src.transform.c) > 360):
                    crs = self._crs_from_epsg_safe(tif_epsg)
                    logger.info('影像坐标系头信息异常，按 EPSG:%s 处理', tif_epsg)

                # CRS 候选列表：文件自带 → 邻近带号 → 用户指定兜底
                # 部分 TIF 头标 EPSG 与实际坐标带号不一致（如青龙闵塆标 4548 实际 4547），
                # 尝试多个候选，选第一个能与边界产生有效交集的
                detected_epsg = crs.to_epsg() if crs is not None else None
                crs_candidates = []
                if crs is not None:
                    crs_candidates.append(('文件自带', crs))
                if detected_epsg in (4547, 4548):
                    neighbor_epsg = 4548 if detected_epsg == 4547 else 4547
                    neighbor_crs = self._crs_from_epsg_safe(neighbor_epsg)
                    crs_candidates.append((f'邻近 EPSG:{neighbor_epsg}', neighbor_crs))
                if tif_epsg and (not detected_epsg or detected_epsg != tif_epsg):
                    crs_candidates.append((f'兜底 EPSG:{tif_epsg}', self._crs_from_epsg_safe(tif_epsg)))

                best_crs = None
                best_bnd_proj = None
                best_win = None
                best_label = ''

                for label, try_crs in crs_candidates:
                    if try_crs is None:
                        continue
                    try:
                        # boundary.to_crs 可能因 PROJ 数据库版本冲突失败，用 Proj4 兜底
                        try:
                            bnd_proj = boundary.to_crs(try_crs)
                        except Exception:
                            logger.info('boundary.to_crs 失败（PROJ 版本冲突），使用 Proj4 兜底变换')
                            from shapely.ops import transform as shapely_transform
                            from pyproj import Transformer
                            target_proj4 = try_crs.to_proj4()
                            transformer = Transformer.from_proj(
                                '+proj=longlat +datum=WGS84 +no_defs', target_proj4, always_xy=True)
                            bnd_geom = boundary.geometry.apply(
                                lambda g: (shapely_transform(transformer.transform, g)
                                           if g and not g.is_empty else g))
                            bnd_proj = gpd.GeoDataFrame(geometry=bnd_geom, crs=try_crs)
                        minx, miny, maxx, maxy = bnd_proj.total_bounds
                        # ── 多片测区扩展：RTK 测点包围盒并入裁剪窗口 ──
                        if points_4326:
                            try:
                                tr_pts = Transformer.from_crs('EPSG:4326', try_crs, always_xy=True)
                            except Exception:
                                tr_pts = None
                            pxs, pys = [], []
                            for lon, lat, _z in points_4326:
                                if tr_pts is not None:
                                    x, y = tr_pts.transform(lon, lat)
                                else:
                                    x, y = lon, lat
                                pxs.append(x)
                                pys.append(y)
                            if pxs:
                                if min(pxs) < minx:
                                    minx = min(pxs)
                                if max(pxs) > maxx:
                                    maxx = max(pxs)
                                if min(pys) < miny:
                                    miny = min(pys)
                                if max(pys) > maxy:
                                    maxy = max(pys)
                                logger.info('测点范围并入裁剪窗口（%d 点）', len(pxs))
                        pad = max(maxx - minx, maxy - miny) * 0.3
                        minx, miny, maxx, maxy = minx - pad, miny - pad, maxx + pad, maxy + pad

                        win = from_bounds(minx, miny, maxx, maxy, transform=src.transform)
                        win = win.intersection(rasterio.windows.Window(0, 0, src.width, src.height))
                        if win.width >= 10 and win.height >= 10:
                            best_crs, best_bnd_proj, best_win, best_label = try_crs, bnd_proj, win, label
                            break
                    except WindowError:
                        continue
                    except Exception:
                        continue

                if best_crs is None:
                    logger.warning('所有 CRS 候选均无法与影像产生交集，尝试的候选: %s',
                                   [c[0] for c in crs_candidates])
                    return None

                if best_label:
                    logger.info('CRS 匹配成功: %s', best_label)
                crs = best_crs
                bnd_proj = best_bnd_proj
                win = best_win
                minx, miny, maxx, maxy = rasterio.windows.bounds(win, src.transform)
                # 限制输出分辨率 ~2000px
                scale = max(win.width, win.height) / 2000.0
                out_w = max(1, int(win.width / max(scale, 1)))
                out_h = max(1, int(win.height / max(scale, 1)))
                count = min(src.count, 3)
                # 调色板影像（单波段+ColorMap，如罗山无人机影像）：索引值不能做
                # average 平均（会得到无效中间色），必须用 nearest 保索引精度，
                # 读取后再应用色表展开为真彩色 RGB（ArcGIS Pro 即如此渲染）。
                try:
                    cmap_check = src.colormap(1)
                except ValueError:
                    cmap_check = None  # 灰度影像无调色板（rasterio 对 NULL color table 抛异常）
                has_cmap = bool(cmap_check) if src.count == 1 else False
                resampling = Resampling.nearest if has_cmap else Resampling.average
                img = src.read(list(range(1, count + 1)), window=win,
                               out_shape=(count, out_h, out_w),
                               resampling=resampling)
                if img.max() == 0:
                    logger.warning('裁剪区域全为空值（影像该处无数据）')
                    return None
                img = np.transpose(img, (1, 2, 0))
                is_gray = False  # 单波段灰度影像标记，后续用 Esri 彩色底图融合
                if img.shape[2] == 1:
                    # 单波段可能携带调色板（ColorMap）：ArcGIS Pro 自动应用色表
                    # 显示彩色，rasterio 默认只给索引值，直接按灰度渲染会丢失颜色。
                    # 有调色板 → 展开为真彩色 RGB；无调色板 → 按灰度拉伸处理。
                    try:
                        cmap = src.colormap(1)
                    except ValueError:
                        cmap = None
                    if cmap:
                        lut = np.zeros((256, 3), dtype=np.uint8)
                        for v, rgb in cmap.items():
                            if 0 <= v < 256:
                                lut[v] = np.asarray(rgb[:3], dtype=np.uint8)
                        idx = np.rint(img[..., 0]).clip(0, 255).astype(np.uint8)
                        img = lut[idx]
                        logger.info('单波段影像已应用 ColorMap 调色板展开为真彩色 RGB（%d 色）', len(cmap))
                    else:
                        is_gray = True
                        # 单波段灰度影像：2%-98% 百分位拉伸增强对比度
                        flat = img[..., 0].ravel()
                        nonzero = flat[flat > 0] if (flat > 0).any() else flat
                        if len(nonzero) >= 10:
                            lo, hi = np.percentile(nonzero, [2, 98])
                        else:
                            lo, hi = flat.min(), flat.max()
                        if hi > lo:
                            stretched = np.clip((img[..., 0].astype(np.float32) - lo) * 255.0 / (hi - lo), 0, 255).astype(np.uint8)
                        else:
                            stretched = img[..., 0]
                        img = np.stack([stretched, stretched, stretched], axis=2)
                        logger.info('单波段影像已做 2%%-98%% 对比度拉伸（%.0f-%.0f → 0-255）', lo, hi)
                # 第4波段为透明度时，用它精准标记无数据区域（供底图填充）
                nodata_mask = None
                if src.count >= 4:
                    alpha = src.read(4, window=win, out_shape=(out_h, out_w),
                                     resampling=Resampling.nearest)
                    nodata_mask = alpha == 0
                win_bounds = rasterio.windows.bounds(win, src.transform)

            extent = (win_bounds[0], win_bounds[2], win_bounds[1], win_bounds[3])
            boundary_xy = [[(c[0], c[1]) for c in geom.exterior.coords] for geom in
                           bnd_proj.geometry.explode(index_parts=False)]

            # 河道管理范围线（hx.geojson 线要素，左岸/右岸分开，全量不处理）
            hedao_xy = []
            hedao = self._load_hedao(boundary)
            if hedao is not None:
                try:
                    hd_proj = hedao.to_crs(crs)
                except Exception:
                    from shapely.ops import transform as shapely_transform
                    from pyproj import Transformer
                    tr = Transformer.from_proj(
                        '+proj=longlat +datum=WGS84 +no_defs', crs.to_proj4(), always_xy=True)
                    hd_geom = hedao.geometry.apply(
                        lambda g: (shapely_transform(tr.transform, g)
                                   if g and not g.is_empty else g))
                    hd_proj = gpd.GeoDataFrame(geometry=hd_geom, crs=crs)
                # 线要素（MultiLineString/LineString）逐线取坐标；面要素兜底取环线
                for geom in hd_proj.geometry:
                    if geom is None or geom.is_empty:
                        continue
                    if geom.geom_type in ('MultiLineString', 'GeometryCollection'):
                        for part in getattr(geom, 'geoms', []):
                            if part.geom_type in ('LineString', 'LinearRing'):
                                hedao_xy.append([(c[0], c[1]) for c in part.coords])
                    elif geom.geom_type in ('LineString', 'LinearRing'):
                        hedao_xy.append([(c[0], c[1]) for c in geom.coords])
                    else:  # Polygon 兜底
                        for ring in [geom.exterior, *geom.interiors]:
                            hedao_xy.append([(c[0], c[1]) for c in ring.coords])
                logger.info('河道管理范围线 %d 条（全量，不裁切）', len(hedao_xy))

            # 底图融合：单波段灰度影像叠加 Esri 彩色卫星底图，消除灰色
            base = self._fetch_esri_basemap(extent, crs, img.shape[0], img.shape[1])
            if base is not None:
                if is_gray:
                    # 单波段灰度影像 → 叠加在 Esri 彩色卫星底图上（50% 透明度），
                    # 既能展示彩色地理背景，又保留无人机影像纹理细节
                    alpha = 0.5
                    img = (base.astype(np.float32) * (1 - alpha) +
                           img.astype(np.float32) * alpha).clip(0, 255).astype(np.uint8)
                    logger.info('单波段灰度影像已融合 Esri 彩色卫星底图（透明度 %.0f%%）', alpha * 100)
                else:
                    nodata = nodata_mask if nodata_mask is not None else np.all(img == 0, axis=2)
                    img[nodata] = base[nodata]
                    logger.info('已垫入 Esri 卫星底图（填充 %.1f%% 无数据区域）',
                                100.0 * nodata.mean())

            plt, fp = self._mpl()
            fig, ax = plt.subplots(figsize=(10, 10 * (extent[3] - extent[2]) /
                                            max(extent[1] - extent[0], 1e-9)), dpi=150)
            ax.imshow(img, extent=extent)
            for coords in hedao_xy:
                xs, ys = zip(*coords)
                ax.plot(xs, ys, color='red', linewidth=2.0)
            for coords in boundary_xy:
                xs, ys = zip(*coords)
                ax.plot(xs, ys, color='yellow', linewidth=2.0)
            ax.set_xlim(extent[0], extent[1])
            ax.set_ylim(extent[2], extent[3])
            ax.axis('off')
            fig.savefig(out_png, bbox_inches='tight', pad_inches=0.05, facecolor='white')
            plt.close(fig)
            logger.info('正射影像裁剪完成: %s', out_png)
            return {'crs': crs, 'extent': extent, 'boundary_xy': boundary_xy,
                    'hedao_xy': hedao_xy, 'img': img}
        except Exception as e:
            logger.exception('影像裁剪失败: %s', e)
            return None

    @staticmethod
    def _fetch_esri_basemap(extent, crs, out_h: int, out_w: int):
        """下载 Esri World Imagery 瓦片并重投影到目标网格（与裁剪影像同尺寸），失败返回 None。"""
        try:
            import math
            from io import BytesIO

            import numpy as np
            import requests
            from PIL import Image
            from pyproj import Transformer
            from rasterio.crs import CRS as RioCRS
            from rasterio.transform import from_bounds
            from rasterio.warp import Resampling, reproject

            # pyproj 可能因 PROJ 数据库冲突无法识别 rasterio CRS
            try:
                tr = Transformer.from_crs(crs, 'EPSG:4326', always_xy=True)
            except Exception:
                tr = Transformer.from_proj(
                    crs.to_proj4(), '+proj=longlat +datum=WGS84 +no_defs', always_xy=True)
            lons, lats = [], []
            for x in (extent[0], extent[1]):
                for y in (extent[2], extent[3]):
                    lon, lat = tr.transform(x, y)
                    lons.append(lon)
                    lats.append(lat)
            lon0, lon1, lat0, lat1 = min(lons), max(lons), min(lats), max(lats)
            lat_c = (lat0 + lat1) / 2
            res_target = (extent[1] - extent[0]) / max(out_w, 1)
            zoom = int(round(math.log2(156543.0339 * math.cos(math.radians(lat_c)) /
                                       max(res_target, 0.01))))
            zoom = max(1, min(18, zoom))

            def tile_xy(lon, lat, z):
                n = 2 ** z
                xt = (lon + 180.0) / 360.0 * n
                yt = (1.0 - math.asinh(math.tan(math.radians(lat))) / math.pi) / 2.0 * n
                return xt, yt

            while True:
                x0f, y0f = tile_xy(lon0, lat1, zoom)  # 左上
                x1f, y1f = tile_xy(lon1, lat0, zoom)  # 右下
                tx0, ty0, tx1, ty1 = int(x0f), int(y0f), int(x1f), int(y1f)
                if (tx1 - tx0 + 1) * (ty1 - ty0 + 1) <= 120 or zoom <= 1:
                    break
                zoom -= 1
            nx, ny = tx1 - tx0 + 1, ty1 - ty0 + 1
            mosaic = np.zeros((ny * 256, nx * 256, 3), dtype=np.uint8)
            sess = requests.Session()
            for ty in range(ty0, ty1 + 1):
                for tx in range(tx0, tx1 + 1):
                    url = ('https://server.arcgisonline.com/ArcGIS/rest/services/'
                           f'World_Imagery/MapServer/tile/{zoom}/{ty}/{tx}')
                    try:
                        r = sess.get(url, timeout=15)
                        if r.status_code != 200:
                            continue
                        tile = np.asarray(Image.open(BytesIO(r.content)).convert('RGB'))
                        mosaic[(ty - ty0) * 256:(ty - ty0 + 1) * 256,
                               (tx - tx0) * 256:(tx - tx0 + 1) * 256] = tile
                    except Exception:
                        continue
            if mosaic.max() == 0:
                return None
            # 马赛克在 3857 下的范围 → 重投影到与裁剪影像一致的网格
            R = 6378137.0
            n = 2 ** zoom

            def merc_x(t):
                return t / n * 2 * math.pi * R - math.pi * R

            def merc_y(t):
                return math.pi * R - t / n * 2 * math.pi * R

            src_transform = from_bounds(merc_x(tx0), merc_y(ty1 + 1),
                                        merc_x(tx1 + 1), merc_y(ty0),
                                        nx * 256, ny * 256)
            dst = np.zeros((3, out_h, out_w), dtype=np.uint8)
            dst_transform = from_bounds(extent[0], extent[2], extent[1], extent[3],
                                        out_w, out_h)
            reproject(source=np.transpose(mosaic, (2, 0, 1)), destination=dst,
                      src_transform=src_transform,
                      src_crs=RioCRS.from_proj4(FALLBACK_PROJ4[3857]),
                      dst_transform=dst_transform, dst_crs=crs,
                      resampling=Resampling.bilinear)
            return np.transpose(dst, (1, 2, 0))
        except Exception as e:
            logger.warning('Esri 底图获取失败（继续无底图出图）: %s', e)
            return None

    def _vlm_interpret(self, image_path: str, site: str, permit_no: str) -> str:
        """qwen-vl 判读：是否停工、河道管理范围内作业器具数量、有无堆砂堆靠。"""
        try:
            import dashscope
            prompt = (f'这是{site}（{permit_no}）采砂区域的无人机正射影像，黄线为批复采区范围，'
                      '红线为河道管理范围线。'
                      '请判读三件事：1）采区是否处于停工状态；2）河道管理范围内可见的采砂船、挖机等'
                      '作业器具数量；3）河道管理范围内有无堆砂、砂堆靠岸堆放现象。'
                      '只依据影像可见内容客观描述，不要推测。用一段正式的报告语言输出（80字以内），'
                      '句式如"经解译，采区内…，河道管理范围内…"。')
            resp = dashscope.MultiModalConversation.call(
                model=VLM_MODEL,
                api_key=os.environ.get('DASHSCOPE_API_KEY'),
                messages=[{'role': 'user', 'content': [
                    {'image': f'file://{image_path}'},
                    {'text': prompt},
                ]}])
            if resp.status_code == 200:
                content = resp.output.choices[0].message.content
                if isinstance(content, list):
                    content = ''.join(c.get('text', '') for c in content)
                return str(content).strip()
            logger.warning('VLM 调用失败: %s', resp.message)
        except Exception as e:
            logger.exception('VLM 解译异常: %s', e)
        return '影像解译未完成（视觉模型调用失败），请人工核查影像。'

    @staticmethod
    def _depth_evaluation(stats: Dict[str, Any]) -> Tuple[str, str]:
        """平均高程 vs 控制高程下限 → 三档结论。返回 (结论文字, 控制高程展示文本)。"""
        ctrl_min, ctrl_max = stats['ctrl_min'], stats['ctrl_max']
        zavg = stats['zavg']
        if ctrl_min is None:
            return (f'实测平均高程{zavg:.3f}m；实施方案中未检索到该采区控制开采高程（{MISSING}），'
                    '无法进行开采深度比对评估。'), ''
        ctrl_text = f'{ctrl_max:.2f}～{ctrl_min:.2f}m'
        diff = ctrl_min - zavg
        if diff > 2:
            verdict = (f'实测平均高程{zavg:.3f}m，低于控制开采高程下限{ctrl_min:.2f}m达{diff:.2f}m（超过2米），'
                       '存在超深度开采。')
        elif diff > 0:
            verdict = (f'实测平均高程{zavg:.3f}m，低于控制开采高程下限{ctrl_min:.2f}m约{diff:.2f}m（2米以内），'
                       '采区局部提砂区修复不到位，后续应加强管理。')
        else:
            verdict = (f'实测平均高程{zavg:.3f}m，不低于控制开采高程下限{ctrl_min:.2f}m，'
                       '本年度开采深度基本符合年度实施方案。')
        return verdict, ctrl_text

    def _draw_elev_figure(self, site: str, points_4326: List[Tuple[float, float, float]],
                          clip: Optional[Dict[str, Any]], out_png: str) -> bool:
        """高程点示意图（仿样例图3.5-29）：视野聚焦测点区域，测线连线 + 红点 + 大号红字标注。"""
        try:
            plt, fp = self._mpl()
            if clip:
                from pyproj import Transformer
                target_crs = clip['crs']
                # pyproj 可能因 PROJ 数据库版本冲突无法识别 rasterio CRS，用 Proj4 兜底
                try:
                    tr = Transformer.from_crs('EPSG:4326', target_crs, always_xy=True)
                except Exception:
                    tr = Transformer.from_proj(
                        '+proj=longlat +datum=WGS84 +no_defs',
                        target_crs.to_proj4(), always_xy=True)
                pts = [(*tr.transform(lon, lat), z) for lon, lat, z in points_4326]
                extent = clip['extent']
                # 视野聚焦到测点包围盒（外扩12%），而非整个裁剪范围
                px = [p[0] for p in pts]
                py = [p[1] for p in pts]
                pad = max(max(px) - min(px), max(py) - min(py)) * 0.12 + 1e-9
                vx0 = max(extent[0], min(px) - pad)
                vx1 = min(extent[1], max(px) + pad)
                vy0 = max(extent[2], min(py) - pad)
                vy1 = min(extent[3], max(py) + pad)
                fig, ax = plt.subplots(figsize=(10, min(16, max(6, 10 * (vy1 - vy0) /
                                                max(vx1 - vx0, 1e-9)))), dpi=150)
                ax.imshow(clip['img'], extent=extent)
                for coords in clip.get('hedao_xy') or []:
                    xs, ys = zip(*coords)
                    ax.plot(xs, ys, color='red', linewidth=1.8)
                for coords in clip['boundary_xy']:
                    xs, ys = zip(*coords)
                    ax.plot(xs, ys, color='lime', linewidth=1.6)
                ax.set_xlim(vx0, vx1)
                ax.set_ylim(vy0, vy1)
            else:
                # 无影像时白底出图
                boundary = self._load_boundary(site)
                pts = points_4326
                fig, ax = plt.subplots(figsize=(10, 8), dpi=150)
                if boundary is not None:
                    for geom in boundary.geometry.explode(index_parts=False):
                        xs = [c[0] for c in geom.exterior.coords]
                        ys = [c[1] for c in geom.exterior.coords]
                        ax.plot(xs, ys, color='green', linewidth=1.8)
            n = len(pts)
            # 注：按用户要求，图内线要素仅保留 2026采区（绿/黄）+ 河道管理范围（红），
            # 不再绘制红色测线连线；高程点以红点+红字标注表达
            import math
            # 与前端 LabelManager 一致的像素空间碰撞检测 + 渲染样式
            # 将数据坐标转为显示像素坐标（同前端 latLngToContainerPoint），
            # 避免按测量序号等间隔跳采样导致 zigzag 测线上相邻轨迹点标签重叠
            import matplotlib.patheffects as pe
            dpi = 150
            # 自适应标签间距：测点密的河段用较大的间距避免视觉混乱，稀疏测点放宽标注密度
            # 基准 65px（适中档），标注太少时逐步降低
            min_label_px = 65   # 标签间最小像素间距（对应前端 LabelManager 适中档位）
            trans = ax.transData
            grid = {}
            label_count = 0
            ms = 3.5 if n <= 500 else 2.0
            LABEL_COLOR = '#ef4444'       # 同前端 Tooltip color
            LABEL_SIZE = 11               # pt，等效前端 14px
            for x, y, z in pts:
                ax.plot(x, y, 'o', color=LABEL_COLOR, markersize=ms, zorder=2)
                px, py = trans.transform((x, y))
                cx = int(px / min_label_px)
                cy = int(py / min_label_px)
                clash = False
                for dx in (-1, 0, 1):
                    for dy in (-1, 0, 1):
                        neighbors = grid.get((cx + dx, cy + dy))
                        if not neighbors:
                            continue
                        for (gx, gy) in neighbors:
                            if math.hypot(px - gx, py - gy) < min_label_px:
                                clash = True
                                break
                        if clash:
                            break
                if not clash:
                    ax.annotate(f'{z:.3f}', (x, y), textcoords='offset points',
                                xytext=(4, 4), fontsize=LABEL_SIZE, color=LABEL_COLOR,
                                fontweight='bold', fontfamily='serif',
                                path_effects=[pe.withStroke(linewidth=3,
                                                            foreground='white')])
                    grid.setdefault((cx, cy), []).append((px, py))
                    label_count += 1
            # 自适应重标注：若首次标注太少（<5个），放宽间距再跑一遍
            if label_count < 5 and n >= 5:
                min_label_px_2 = max(30, min_label_px // 2)
                logger.info('标注过少（%d/%d），降低碰撞阈值 %d→%d 重试', label_count, n, min_label_px, min_label_px_2)
                # 清除已绘制的文本标注（annotations）但保留红点和连线
                for child in ax.get_children():
                    if isinstance(child, type(ax.annotate('', (0, 0)))):
                        child.remove()
                grid.clear()
                label_count = 0
                for x, y, z in pts:
                    px, py = trans.transform((x, y))
                    cx = int(px / min_label_px_2)
                    cy = int(py / min_label_px_2)
                    clash = False
                    for dx in (-1, 0, 1):
                        for dy in (-1, 0, 1):
                            neighbors = grid.get((cx + dx, cy + dy))
                            if not neighbors:
                                continue
                            for (gx, gy) in neighbors:
                                if math.hypot(px - gx, py - gy) < min_label_px_2:
                                    clash = True
                                    break
                            if clash:
                                break
                    if not clash:
                        ax.annotate(f'{z:.3f}', (x, y), textcoords='offset points',
                                    xytext=(4, 4), fontsize=LABEL_SIZE, color=LABEL_COLOR,
                                    fontweight='bold', fontfamily='serif',
                                    path_effects=[pe.withStroke(linewidth=3,
                                                                foreground='white')])
                        grid.setdefault((cx, cy), []).append((px, py))
                        label_count += 1
            ax.axis('off')
            fig.savefig(out_png, bbox_inches='tight', pad_inches=0.05, facecolor='white')
            plt.close(fig)
            logger.info('高程点示意图完成: %s（%d 点，空间去重后标注 %d 个）',
                        out_png, n, label_count)
            return True
        except Exception as e:
            logger.exception('高程点示意图失败: %s', e)
            return False

    @staticmethod
    def _mpl():
        from tools.report_generator_tool import ReportGeneratorTool
        return ReportGeneratorTool()._setup_matplotlib_chinese()

    # ------------------------------------------------------------------
    # docx 渲染
    # ------------------------------------------------------------------
    def _render_docx(self, variables: Dict[str, str], ortho_png: Optional[str],
                     elev_png: Optional[str],
                     field_pngs: Optional[List[str]] = None) -> Dict[str, Any]:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt
        from docx.text.paragraph import Paragraph
        from tools.report_generator_tool import ReportGeneratorTool

        template_path = os.path.join(BACKEND_DIR, 'templates', 'caisha_monitor_template.docx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(f'模板不存在: {template_path}，请先运行 create_caisha_template.py')

        doc = Document(template_path)
        site = variables.get('site_name', '')
        display_name = variables.get('display_name') or site
        field_pngs = [p for p in (field_pngs or []) if p and os.path.exists(p)]

        def _insert_after(anchor):
            new_p = OxmlElement('w:p')
            anchor._p.addnext(new_p)
            p = Paragraph(new_p, anchor._parent)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return p

        def _caption_run(p, text):
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        # 图片占位段落 → 原位插图（先处理，避免被文本替换吃掉）
        img_map = {'{{img_drone_ortho}}': ortho_png, '{{img_elev_points}}': elev_png}
        for paragraph in list(doc.paragraphs):
            text = paragraph.text.strip()
            if text == '{{img_field_photos}}':
                # 现场监测图片：数量不定，逐张插图+图名；无图时清空占位段落
                for run in paragraph.runs:
                    run.text = ''
                cursor = paragraph
                for idx, img in enumerate(field_pngs):
                    if idx == 0:
                        pic_p = paragraph
                    else:
                        pic_p = _insert_after(cursor)
                    pic_p.add_run().add_picture(img, width=Inches(5.0))
                    cap_p = _insert_after(pic_p)
                    suffix = f'（{"一二三四五六"[idx]}）' if len(field_pngs) > 1 else ''
                    _caption_run(cap_p, f'图{idx + 1} {display_name}现场监测图片{suffix}')
                    cursor = cap_p
            elif text in img_map:
                for run in paragraph.runs:
                    run.text = ''
                img = img_map[text]
                if img and os.path.exists(img):
                    paragraph.add_run().add_picture(img, width=Inches(5.8))
                else:
                    paragraph.add_run('（影像图缺失）')

        helper = ReportGeneratorTool()
        for paragraph in doc.paragraphs:
            helper._replace_text_in_paragraph(paragraph, variables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        helper._replace_text_in_paragraph(paragraph, variables)

        out_dir = os.path.join(BACKEND_DIR, 'static', 'reports')
        os.makedirs(out_dir, exist_ok=True)
        # 秒级时间戳在快速连续生成时会撞名覆盖，加随机后缀保证唯一
        import uuid
        filename = f"caisha_report_{int(time.time())}_{uuid.uuid4().hex[:8]}.docx"
        output_path = os.path.join(out_dir, filename)
        doc.save(output_path)
        return {
            'success': True,
            'message': f"{variables.get('site_name', '')}采砂监测分析报告生成成功",
            'file_path': output_path,
            'download_url': f'/api/download/report/{filename}',
            'filename': filename,
        }
