import os
import re
import time
import json
import logging
import glob
import tempfile
import requests
from datetime import datetime
from typing import Dict, Any, Union, List
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from qwen_agent.tools.base import BaseTool, register_tool

logger = logging.getLogger(__name__)

SCREENSHOT_DIR = "/home/server/python/map_assistant_v1/backend/static/screenshots"
KB_IMAGE_BASE_URL = "http://172.136.16.52:82"

@register_tool('report_generator_tool')
class ReportGeneratorTool(BaseTool):
    description = '''
    报告生成工具（自动知识库集成版）。用于根据提供的变量数据，基于固定模板生成 Word 文档报告。
    
    【自动知识库检索】
    ⚠️ 重要：本工具会在生成报告时自动检索知识库获取相关政策、规范、流程等信息！
    - 如果 variables 中没有 knowledge_content 字段，工具会自动推断检索关键词并查询知识库
    - 检索结果会自动添加到 variables 的 knowledge_content 字段中
    - 建议在模板中使用 {{knowledge_content}} 占位符来展示知识库内容
    
    使用场景：当用户请求生成报告、文档或总结时使用。
    
    使用步骤（已自动化，无需手动调用知识库）：
    1. 调用 postgresql_tool 获取业务数据
    2. 调用本工具（会自动检索知识库）
    3. 本工具会自动整合业务数据 + 知识库内容
    
    注意：generated_date 会自动填充当前日期，无需手动传入。
    若有地图截图路径（map_image_path），传入后会自动嵌入报告。
    
    【模板占位符建议】
    - report_title: 报告标题
    - summary: 摘要（会自动整合知识库内容）
    - details: 详细信息
    - knowledge_content: 知识库检索内容（自动填充）
    - knowledge_references: 知识库引用说明
    - conclusion: 结论
    - recommendations: 建议
    '''
    
    parameters = [
        {
            'name': 'variables',
            'type': 'object',
            'description': '用于替换模板中占位符的变量字典。Key为占位符名称（不含花括号），Value为替换文本。例如：{"report_title": "xx分析报告", "summary": "..."}',
            'required': True
        },
        {
            'name': 'map_image_path',
            'type': 'string',
            'description': '地图截图的服务器绝对路径（如 /home/.../static/screenshots/xxx.png）。若传入，会在报告正文末尾自动插入地图截图。',
            'required': False
        },
        {
            'name': 'template_name',
            'type': 'string',
            'description': '模板文件名（默认为 report_template.docx）',
            'required': False
        }
    ]

    def call(self, params: Union[str, Dict[str, Any]], **kwargs) -> Dict[str, Any]:
        if isinstance(params, str):
            try:
                params = json.loads(params)
            except json.JSONDecodeError:
                return {'success': False, 'error': 'Invalid JSON parameters'}

        variables = params.get('variables', {})
        template_name = params.get('template_name', 'report_template.docx')
        map_image_path = params.get('map_image_path', None)
        
        # 强制检查：report_generator_tool 必须在 knowledge_base_tool 之后调用
        # 如果 variables 中没有知识库内容，自动检索
        if not variables.get('knowledge_content') and not variables.get('knowledge_references'):
            logger.warning('⚠️ 警告：report_generator_tool 被调用时未提供知识库内容，将自动检索...')
            
            # 从 report_title 推断检索关键词
            report_title = variables.get('report_title', '')
            search_keywords = self._infer_search_keywords(report_title, variables)
            
            if search_keywords:
                logger.info(f'🔍 自动检索知识库: {search_keywords}')
                kb_result = self._auto_search_knowledge(search_keywords)
                if kb_result.get('success'):
                    kb_content = kb_result.get('content', '')
                    kb_images = kb_result.get('images', [])  # 提取图片路径
                    variables['knowledge_content'] = kb_content
                    variables['knowledge_references'] = f'知识库检索结果（关键词：{search_keywords}）'
                    # 保存原始图片路径列表，供后续下载插入 Word
                    variables['_kb_image_paths'] = kb_images
                    if kb_images:
                        logger.info(f'✅ 知识库自动检索成功: 文本{len(kb_content)}字, 图片{len(kb_images)}张')
                    else:
                        logger.info(f'✅ 知识库自动检索成功，获取内容长度: {len(kb_content)} 字符')
                    
                    # 打印知识库检索到的实际内容（方便调试）
                    logger.info(f'📄 [KB内容] {kb_content}')
                else:
                    logger.error(f'❌ 知识库自动检索失败: {kb_result.get("error")}')
                    variables['knowledge_content'] = ''
                    variables['knowledge_references'] = '知识库检索失败'
        else:
            logger.info('✅ 检测到已提供知识库内容，跳过自动检索')
        
        # 无条件强制覆盖 generated_date，确保始终使用服务器当前真实时间
        variables['generated_date'] = datetime.now().strftime('%Y年%m月%d日')
        
        # 清理所有变量值中的 Markdown 格式符号，写入 Word 使用纯净文字
        for key in list(variables.keys()):
            if isinstance(variables[key], str):
                variables[key] = self._strip_markdown(variables[key])
        
        # 路径配置
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_path = os.path.join(base_dir, 'templates', template_name)
        output_dir = os.path.join(base_dir, 'static', 'reports')
        
        if not os.path.exists(template_path):
            return {'success': False, 'error': f'Template {template_name} not found at {template_path}'}
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        try:
            doc = Document(template_path)
            
            # 替换段落中的文本（修复跨 run 占位符）
            for paragraph in doc.paragraphs:
                self._replace_text_in_paragraph(paragraph, variables)
            
            # 替换表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_text_in_paragraph(paragraph, variables)

            # 插入自动统计图表到"详细数据"段落之后
            report_charts = variables.get('_report_charts', [])
            if isinstance(report_charts, list) and report_charts:
                self._insert_generated_charts(doc, report_charts, after_section='details')

            # 插入知识库图片（从远程 URL 下载后嵌入 Word）
            kb_image_paths = variables.get('_kb_image_paths', [])
            if kb_image_paths:
                self._insert_kb_images(doc, kb_image_paths)

            # 插入地图截图
            img_path = self._resolve_image_path(map_image_path)
            if img_path and os.path.exists(img_path):
                try:
                    doc.add_paragraph('地图截图：', style='Heading 2' if 'Heading 2' in [s.name for s in doc.styles] else None)
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(6.0))
                    logger.info(f'已插入地图截图: {img_path}')
                except Exception as img_err:
                    logger.warning(f'插入截图失败: {img_err}')
            else:
                logger.warning(f'未找到有效截图路径: {map_image_path}')

            # 生成唯一文件名
            timestamp = int(time.time())
            filename = f"report_{timestamp}.docx"
            output_path = os.path.join(output_dir, filename)
            
            doc.save(output_path)
            
            download_url = f"/api/download/report/{filename}"
            
            return {
                'success': True,
                'message': '报告生成成功',
                'file_path': output_path,
                'download_url': download_url,
                'filename': filename
            }
            
        except Exception as e:
            logger.error(f"Error generating report: {e}")
            return {'success': False, 'error': str(e)}

    def _strip_markdown(self, text: str) -> str:
        """去除 Markdown 格式符，输出纯净可读文字"""
        # 去除标题符号 ### ## #
        text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
        # 去除加粗/斜体 **xx** / *xx* / __xx__ / _xx_
        text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', text)
        text = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', text)
        # 去除行内代码 `xx`
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # 去除链接 [text](url)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # 去除无序列表符号 - / * 开头
        text = re.sub(r'^[\-\*]\s+', '', text, flags=re.MULTILINE)
        # 去除有序列表序号 1. 2.
        text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE)
        # 压缩多余空行（超过2个换行合并为2个）
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _resolve_image_path(self, map_image_path):
        """解析截图路径：优先用传入路径，否则自动取最新截图"""
        if map_image_path and os.path.exists(map_image_path):
            return map_image_path
        # 自动取最新截图
        pattern = os.path.join(SCREENSHOT_DIR, '*.png')
        files = sorted(glob.glob(pattern), key=os.path.getmtime, reverse=True)
        if files:
            logger.info(f'自动使用最新截图: {files[0]}')
            return files[0]
        return None

    def _build_image_url(self, raw_path: str) -> str:
        """将知识库返回的相对图片路径拼接为完整可下载 URL。"""
        path = raw_path.strip()
        # 已经是完整 URL
        if path.startswith('http://') or path.startswith('https://'):
            return path
        # 去除开头多余的 /
        path = path.lstrip('/')
        return f"{KB_IMAGE_BASE_URL}/{path}"

    def _download_image(self, url: str, timeout: int = 15) -> str:
        """下载远程图片到临时文件，返回本地路径；失败返回空字符串。"""
        try:
            resp = requests.get(url, timeout=timeout, stream=True)
            resp.raise_for_status()
            # 从 URL 推断扩展名
            ext = os.path.splitext(url.split('?')[0])[-1] or '.jpg'
            if ext.lower() not in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'):
                ext = '.jpg'
            tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
            for chunk in resp.iter_content(8192):
                tmp.write(chunk)
            tmp.close()
            logger.info(f'[KB Image] 下载成功: {url} -> {tmp.name}')
            return tmp.name
        except Exception as e:
            logger.warning(f'[KB Image] 下载失败: {url}, 错误: {e}')
            return ''

    def _insert_kb_images(self, doc: Document, image_paths: List[str]):
        """将知识库图片下载并插入 Word 文档末尾。"""
        inserted = 0
        temp_files = []
        try:
            doc.add_paragraph('相关影像资料：', style='Heading 2' if 'Heading 2' in [s.name for s in doc.styles] else None)
            for raw_path in image_paths:
                url = self._build_image_url(raw_path)
                local_path = self._download_image(url)
                if not local_path:
                    continue
                temp_files.append(local_path)
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(local_path, width=Inches(5.5))
                    inserted += 1
                    logger.info(f'[KB Image] 已插入第 {inserted} 张: {raw_path}')
                except Exception as e:
                    logger.warning(f'[KB Image] 插入 Word 失败: {raw_path}, 错误: {e}')
            logger.info(f'[KB Image] 共插入 {inserted}/{len(image_paths)} 张知识库图片')
        finally:
            # 清理临时文件
            for tf in temp_files:
                try:
                    os.unlink(tf)
                except OSError:
                    pass

    _cn_font_prop = None

    def _setup_matplotlib_chinese(self):
        """配置 matplotlib 使用系统中文字体，返回 (plt, FontProperties)。"""
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        from matplotlib import font_manager as fm

        plt.rcParams['axes.unicode_minus'] = False

        if ReportGeneratorTool._cn_font_prop is not None:
            return plt, ReportGeneratorTool._cn_font_prop

        candidates = [
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc',
            '/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc',
        ]
        font_path = next((p for p in candidates if os.path.exists(p)), None)

        if not font_path:
            # 用 fc-list 兜底搜一个 CJK 字体
            try:
                import subprocess
                out = subprocess.run(['fc-list', ':lang=zh', '-f', '%{file}\n'],
                                     capture_output=True, text=True).stdout
                for line in out.splitlines():
                    line = line.strip()
                    if line and os.path.exists(line):
                        font_path = line
                        break
            except Exception:
                pass

        if font_path:
            fm.fontManager.addfont(font_path)
            prop = fm.FontProperties(fname=font_path)
            plt.rcParams['font.family'] = prop.get_name()
            ReportGeneratorTool._cn_font_prop = prop
            logger.info(f'[Chart] 中文字体已加载: {font_path} ({prop.get_name()})')
            return plt, prop

        logger.warning('[Chart] 未找到 CJK 字体，中文可能显示为方块')
        return plt, None

    def _draw_bar_chart(self, title: str, labels: List[str], values: List[float], unit: str = '') -> str:
        """绘制水平柱状图并返回临时图片路径。"""
        plt, fp = self._setup_matplotlib_chinese()

        fig, ax = plt.subplots(figsize=(7.5, max(3.0, len(values) * 0.9)), dpi=150)
        colors = ['#3b82f6', '#8b5cf6', '#06b6d4', '#f59e0b', '#ef4444', '#22c55e']
        y_pos = list(range(len(labels)))
        bars = ax.barh(y_pos, values, color=colors[:len(values)], height=0.55,
                        edgecolor='white', linewidth=0.5)

        ax.set_yticks(y_pos)
        ax.set_yticklabels(labels, fontsize=11, fontproperties=fp)
        ax.invert_yaxis()
        ax.set_title(title, fontsize=13, fontweight='bold', pad=12, fontproperties=fp)
        ax.set_xlabel(f'数值（{unit}）' if unit else '数值', fontsize=10, fontproperties=fp)
        ax.grid(axis='x', linestyle='--', alpha=0.3)
        ax.set_axisbelow(True)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        for tick in ax.get_xticklabels():
            if fp is not None:
                tick.set_fontproperties(fp)

        max_v = max(values) if values else 1
        for b, v in zip(bars, values):
            ax.text(
                b.get_width() + max_v * 0.02,
                b.get_y() + b.get_height() / 2,
                f"{v:.3f} {unit}".strip(),
                ha='left', va='center', fontsize=9, color='#334155',
                fontproperties=fp,
            )

        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        fig.tight_layout()
        fig.savefig(tmp.name, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        return tmp.name

    def _draw_doughnut_chart(self, title: str, labels: List[str], values: List[float], note: str = '') -> str:
        """绘制环形图并返回临时图片路径。"""
        plt, fp = self._setup_matplotlib_chinese()

        fig, ax = plt.subplots(figsize=(6.0, 4.5), dpi=150)
        colors = ['#ef4444', '#22c55e', '#3b82f6', '#f59e0b']
        wedges, texts, autotexts = ax.pie(
            values,
            labels=labels,
            colors=colors[:len(values)],
            autopct=lambda p: f'{p:.1f}%' if p > 0.5 else '',
            startangle=90,
            wedgeprops={'width': 0.42, 'edgecolor': 'white', 'linewidth': 2},
            textprops={'fontsize': 10},
            pctdistance=0.78,
        )
        for t in texts:
            if fp is not None:
                t.set_fontproperties(fp)
        for at in autotexts:
            at.set_fontsize(10)
            at.set_fontweight('bold')
            at.set_color('white')
            if fp is not None:
                at.set_fontproperties(fp)

        ax.set_title(title, fontsize=13, fontweight='bold', pad=10, fontproperties=fp)
        ax.axis('equal')

        if note:
            fig.text(0.5, 0.02, note, ha='center', va='center',
                     fontsize=10, color='#475569', fontproperties=fp)

        tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        fig.tight_layout(rect=[0, 0.05, 1, 1])
        fig.savefig(tmp.name, bbox_inches='tight', facecolor='white')
        plt.close(fig)
        return tmp.name

    def _find_details_paragraph_index(self, doc: Document) -> int:
        """找到"详细数据"段落（{{details}} 替换后的内容）的最后一个段落索引。"""
        body = doc.element.body
        paragraphs = doc.paragraphs
        details_idx = -1

        for i, p in enumerate(paragraphs):
            text = p.text.strip()
            if '详细数据' in text or '{{details}}' in text:
                details_idx = i

        if details_idx < 0:
            return -1

        # 往后扫描：跳过紧跟的非空文本段落（属于 details 替换后的多段内容）
        last_idx = details_idx
        for j in range(details_idx + 1, len(paragraphs)):
            text = paragraphs[j].text.strip()
            # 遇到下一个标题（如"3. 结论"）或空白行则停止
            if not text or re.match(r'^\d+[.、．]', text):
                break
            last_idx = j

        return last_idx

    def _insert_generated_charts(self, doc: Document, charts: List[Dict[str, Any]], after_section: str = 'details'):
        """根据结构化数据自动生成并插入图表到指定章节之后。"""
        from copy import deepcopy

        inserted = 0
        temp_files: List[str] = []

        try:
            # 定位插入锚点
            anchor_idx = self._find_details_paragraph_index(doc)
            body = doc.element.body

            if anchor_idx >= 0:
                anchor_elem = doc.paragraphs[anchor_idx]._element
            else:
                anchor_elem = None
                logger.warning('[Report Chart] 未找到详细数据段落，图表将追加到文档末尾')

            chart_elements = []

            for item in charts:
                if not isinstance(item, dict):
                    continue

                chart_type = str(item.get('type', '')).lower()
                title = str(item.get('title', ''))
                labels = item.get('labels', []) or []
                values = item.get('values', []) or []
                unit = str(item.get('unit', '') or '')
                note = str(item.get('note', '') or '')

                try:
                    num_values = [float(v) for v in values]
                except Exception:
                    logger.warning(f'[Report Chart] skip invalid: {item}')
                    continue

                if not labels or not num_values or len(labels) != len(num_values):
                    continue

                img_path = ''
                if chart_type == 'bar':
                    img_path = self._draw_bar_chart(title, labels, num_values, unit)
                elif chart_type in ('doughnut', 'pie'):
                    img_path = self._draw_doughnut_chart(title, labels, num_values, note)
                else:
                    continue

                if not img_path:
                    continue

                temp_files.append(img_path)
                p = doc.add_paragraph()
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                chart_elements.append(p._element)
                inserted += 1
                logger.info(f'[Report Chart] generated #{inserted}: {title}')

            # 将图表段落从文档末尾移动到锚点之后
            if anchor_elem is not None and chart_elements:
                for elem in chart_elements:
                    body.remove(elem)
                insert_after = anchor_elem
                for elem in chart_elements:
                    insert_after.addnext(elem)
                    insert_after = elem
                logger.info(f'[Report Chart] {inserted} charts inserted after details section')

        except Exception as e:
            logger.warning(f'[Report Chart] error: {e}')
        finally:
            for tf in temp_files:
                try:
                    os.unlink(tf)
                except OSError:
                    pass

    def _replace_text_in_paragraph(self, paragraph, variables):
        """
        修复跨 run 占位符替换。
        - 若替换值中含有 \\n\\n，则把当前段落拆分为多个段落（保持原始样式）。
        - 若替换值中含有 \\n（单行），用软回车（XML <w:br/>）分隔。
        """
        full_text = ''.join(run.text for run in paragraph.runs)
        if '{{' not in full_text:
            return

        # 找到需要多段落展开的占位符
        for key, value in variables.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder not in full_text:
                continue

            value_str = str(value)

            # 含有双换行 → 拆为多个段落
            if '\n\n' in value_str:
                segments = [s.strip() for s in value_str.split('\n\n') if s.strip()]
                if len(segments) > 1:
                    self._replace_with_paragraphs(paragraph, placeholder, segments)
                    return  # 已处理，跳出（段落引用已失效）
                else:
                    value_str = segments[0] if segments else ''

            # 含有单换行 → 替换为段落内换行（<w:br/>）
            if '\n' in value_str:
                self._replace_with_linebreaks(paragraph, placeholder, value_str)
                return

            # 普通替换
            full_text = full_text.replace(placeholder, value_str)

        replaced = full_text
        # 写回
        if paragraph.runs:
            paragraph.runs[0].text = replaced
            for run in paragraph.runs[1:]:
                run.text = ''

    def _replace_with_paragraphs(self, anchor_paragraph, placeholder: str, segments: List[str]):
        """
        将占位符所在段落替换为多个段落，每个 segment 对应一个段落。
        保留原始段落的样式（style）和字体格式。
        """
        from docx.oxml.ns import qn as _qn
        from copy import deepcopy
        import lxml.etree as _etree

        parent = anchor_paragraph._element.getparent()
        idx = list(parent).index(anchor_paragraph._element)

        # 取原始段落的 pPr（段落格式）作为模板
        original_pPr = anchor_paragraph._element.find(_qn('w:pPr'))
        # 取原始 run 的 rPr（字符格式）
        original_rPr = None
        if anchor_paragraph.runs:
            original_rPr = anchor_paragraph.runs[0]._element.find(_qn('w:rPr'))

        # 删除原始占位符段落
        parent.remove(anchor_paragraph._element)

        # 逐段插入
        for i, segment in enumerate(segments):
            new_p = OxmlElement('w:p')

            # 复制段落格式
            if original_pPr is not None:
                new_p.append(deepcopy(original_pPr))

            # 添加 run
            new_r = OxmlElement('w:r')
            if original_rPr is not None:
                new_r.append(deepcopy(original_rPr))
            new_t = OxmlElement('w:t')
            new_t.text = segment
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_r.append(new_t)
            new_p.append(new_r)

            parent.insert(idx + i, new_p)

    def _replace_with_linebreaks(self, paragraph, placeholder: str, value_str: str):
        """
        将段落内的占位符替换为带软回车的内容（同一段落内换行）。
        """
        lines = value_str.split('\n')
        if not paragraph.runs:
            return

        # 清空所有 run
        first_run = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run.text = ''
        first_run.text = ''

        # 在第一个 run 里逐行插入文字 + <w:br/>
        r_elem = first_run._element
        # 清除 r_elem 下现有 w:t
        for t in r_elem.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
            r_elem.remove(t)

        for i, line in enumerate(lines):
            t = OxmlElement('w:t')
            t.text = line
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_elem.append(t)
            if i < len(lines) - 1:
                br = OxmlElement('w:br')
                r_elem.append(br)
    
    def _infer_search_keywords(self, report_title: str, variables: dict) -> str:
        """
        从报告标题和变量内容推断知识库检索关键词。
        策略：
        1. 始终以 report_title 作为核心检索词（最具体，优先级最高）
        2. 从标题中提取补充主题词（如"超深度开采"、"测深监测"等）
        3. 最终拼接为："{report_title} {补充词}"，确保检索能命中具体名称
        """
        # 清理标题中的括号内容（如年份、日期），保留核心词
        core_title = re.sub(r'[（(][^）)]*[）)]', '', report_title).strip()
        core_title = re.sub(r'报告$', '', core_title).strip()

        # 补充主题词映射（只从标题提取，不做宽泛匹配）
        supplement_map = [
            ('超深度开采', '超深度开采'),
            ('超深', '超深度开采'),
            ('测深', '测深监测'),
            ('监测评估', '监测评估'),
            ('无人机', '无人机航测'),
            ('水下测量', '水下测量'),
        ]

        supplements = []
        for pattern, word in supplement_map:
            if pattern in report_title and word not in supplements:
                supplements.append(word)

        # 核心词（具体采区/砂场名称）始终放在最前面
        parts = [core_title] + supplements
        query = ' '.join(parts[:3])

        logger.info(f'🔍 知识库检索词（基于标题）: {query}')
        return query

    def _extract_sand_site_name(self, text: str) -> str:
        """从检索词中提取采砂场/可采区名称，用于元数据约束。"""
        if not text:
            return ''
        m = re.search(r'([\u4e00-\u9fa5A-Za-z0-9\-]{2,40}(?:可采区|采砂场|砂场))', text)
        return m.group(1).strip() if m else ''

    def _flatten_meta_values(self, value) -> List[str]:
        """将元数据值拍平成字符串列表。"""
        out = []
        if value is None:
            return out
        if isinstance(value, (str, int, float, bool)):
            s = str(value).strip()
            if s:
                out.append(s)
            return out
        if isinstance(value, list):
            for item in value:
                out.extend(self._flatten_meta_values(item))
            return out
        if isinstance(value, dict):
            for _, v in value.items():
                out.extend(self._flatten_meta_values(v))
            return out
        s = str(value).strip()
        if s:
            out.append(s)
        return out

    def _filter_chunks_by_sand_site(self, chunks: List[Dict[str, Any]], target_site: str) -> List[Dict[str, Any]]:
        """按元数据中的采砂场名称/可采区名称过滤切片。"""
        if not target_site:
            return chunks

        site_fields = ('采砂场名称', '可采区名称', '采区名称', '采砂场', '可采区', '作业地点')
        filtered = []

        for chunk in chunks:
            meta = chunk.get('meta_fields') or chunk.get('metadata') or {}
            candidates = []

            if isinstance(meta, dict):
                for k, v in meta.items():
                    key = str(k)
                    values = self._flatten_meta_values(v)
                    if any(f in key for f in site_fields):
                        candidates.extend(values)

                if not candidates:
                    candidates.extend(self._flatten_meta_values(meta))
            else:
                candidates.extend(self._flatten_meta_values(meta))

            if any((target_site in c) or (c in target_site) for c in candidates if c):
                filtered.append(chunk)

        return filtered
    
    def _auto_search_knowledge(self, keywords: str) -> dict:
        """
        自动执行知识库检索
        尝试导入并使用 RagFlow 知识库工具
        """
        try:
            # 尝试导入 RagFlow 知识库工具
            import requests
            
            api_key = 'ragflow-jZ-6x-X_PGr5ULHFSPqWhfbmd-0xlU_naoGg0hLc3K0'
            api_base = 'http://172.136.16.14:8080/api/v1'
            dataset_id = '538b0a5c36ff11f18e7d3d43671e73e4'
            
            # 调用检索 API
            url = f"{api_base}/retrieval"
            headers = {
                'Authorization': f'Bearer {api_key}',
                'Content-Type': 'application/json'
            }
            
            payload = {
                'question': keywords,
                'dataset_ids': [dataset_id],
                'top_k': 50,            # 与 RagFlow 测试界面对齐
                'similarity_threshold': 0.2,
                'vector_similarity_weight': 0.5,
                'rerank_id': 'gte-rerank'  # 启用 Rerank 模型排序
            }
            
            response = requests.post(url, headers=headers, json=payload, timeout=15)
            response.raise_for_status()
            
            # 安全解析 JSON，防止非 JSON 响应导致崩溃
            try:
                data = response.json()
            except Exception as json_err:
                logger.error(f'[KB] RagFlow API 返回非 JSON 响应: {response.text[:500]}')
                return {'success': False, 'error': f'知识库检索失败: API 返回非 JSON 数据'}
            
            if not isinstance(data, dict):
                logger.error(f'[KB] RagFlow API 返回非字典类型: {type(data)}')
                return {'success': False, 'error': '知识库检索失败: API 返回格式异常'}
            
            if data.get('code') != 0:
                return {'success': False, 'error': f"知识库检索失败: {data.get('message', '未知错误')}"}
            
            # 提取检索结果
            chunks = data.get('data', {}).get('chunks', []) if isinstance(data.get('data'), dict) else []
            
            # 调试：打印返回数据结构，确认图片路径是否在 API 响应中
            if chunks:
                first_chunk_keys = list(chunks[0].keys()) if isinstance(chunks[0], dict) else 'not dict'
                first_chunk_preview = str(chunks[0].get('content', ''))[:200] if isinstance(chunks[0], dict) else str(chunks[0])[:200]
                logger.info(f'[KB] API 返回 {len(chunks)} 条切片, 字段: {first_chunk_keys}')
                logger.info(f'[KB] 第1条预览: {first_chunk_preview}')
                # 打印前3条完整内容用于对比 RagFlow 测试界面
                for i, c in enumerate(chunks[:3]):
                    content = c.get('content', '') if isinstance(c, dict) else ''
                    has_img = bool(re.search(r'\(images/[^)]+\.(jpg|png|jpeg)\)', content, re.I))
                    has_relevant = '【相关图片】' in content or '【相关影像】' in content
                    logger.info(f'[KB] 切片#{i+1} ({len(content)}字): has_images={has_img}, has_相关图片={has_relevant}')
                    # 如果有图片相关内容，打印完整内容
                    if has_img or has_relevant:
                        logger.info(f'[KB] >>> 切片#{i+1} 完整内容:\n{content}')
            if not chunks:
                return {'success': False, 'error': '未找到相关知识'}

            # 先按“采砂场名称”元数据做约束，避免混入其他可采区的图文
            target_site = self._extract_sand_site_name(keywords)
            candidate_chunks = chunks
            if target_site:
                by_site = self._filter_chunks_by_sand_site(chunks, target_site)
                if by_site:
                    candidate_chunks = by_site
                    logger.info(f'[KB] 采砂场名称约束生效: {target_site}, 命中 {len(by_site)}/{len(chunks)} 条切片')
                else:
                    logger.warning(f'[KB] 采砂场名称约束未命中: {target_site}, 回退使用原始检索结果')
            
            # === 策略：前2条用于文本内容；若第3条是第2条续块，则合并到第2条 ===
            # 同时仅从选中的文本块中提取图片，避免把不相关切片图片带入报告
            all_cleaned_parts = []
            all_images = []
            img_pattern = re.compile(r'[（(]([^)）]*?\.(?:jpg|png|jpeg|gif|bmp))[）)]', re.IGNORECASE)

            selected_chunks = list(candidate_chunks[:2])
            merged_chunk3 = False

            if len(candidate_chunks) >= 3 and len(selected_chunks) >= 2:
                second = selected_chunks[1]
                third = candidate_chunks[2]

                second_doc = str(second.get('document_id', '') or second.get('doc_id', '')).strip()
                third_doc = str(third.get('document_id', '') or third.get('doc_id', '')).strip()
                second_title = str(second.get('docnm_kwd', '') or second.get('title', '')).strip()
                third_title = str(third.get('docnm_kwd', '') or third.get('title', '')).strip()
                third_raw = (third.get('content') or '').strip()

                same_doc = (second_doc and third_doc and second_doc == third_doc) or (
                    second_title and third_title and second_title == third_title
                )
                third_has_images = bool(img_pattern.search(third_raw)) or ('【相关图片】' in third_raw or '【相关影像】' in third_raw)
                third_short = len(third_raw) <= 280

                if same_doc and (third_has_images or third_short):
                    merged = dict(second)
                    merged['content'] = f"{(second.get('content') or '').rstrip()}\n{third_raw}".strip()
                    selected_chunks[1] = merged
                    merged_chunk3 = True
                    logger.info('[KB] 检测到第3条为第2条续块，已执行 2+3 合并策略。')
            
            # 1) 处理选中的文本切片（前2条，必要时第3条并入第2条）
            for rank, chunk in enumerate(selected_chunks, start=1):
                raw = chunk.get('content', '').strip()
                if not raw:
                    continue
                
                score = round(chunk.get('similarity', 0), 3)
                cleaned, extracted_images = self._clean_chunk(raw)
                
                if cleaned:
                    all_cleaned_parts.append(cleaned)
                    all_images.extend(extracted_images)
                    logger.info(f'[KB] 切片#{rank} 文本: {len(raw)}→{len(cleaned)}字, 分数={score}, 图片{len(extracted_images)}张')
            
            # 2) 仅扫描选中切片提取图片，避免引入不相关切片图片
            for rank, chunk in enumerate(selected_chunks, start=1):
                raw = chunk.get('content', '')
                if not raw:
                    continue
                found = img_pattern.findall(raw)
                if found:
                    for p in found:
                        clean_p = p.strip()
                        if clean_p and clean_p not in all_images:
                            all_images.append(clean_p)
                            logger.info(f'[KB] 从切片#{rank} 发现图片: {clean_p}')
            
            # 3) 检查【相关图片】块格式（RagFlow 标准格式，使用中文括号）
            for rank, chunk in enumerate(selected_chunks, start=1):
                raw = chunk.get('content', '')
                if '【相关图片】' in raw or '【相关影像】' in raw:
                    block_imgs = img_pattern.findall(raw)
                    for p in block_imgs:
                        clean_p = p.strip()
                        if clean_p and clean_p not in all_images:
                            all_images.append(clean_p)
                            logger.info(f'[KB] 从切片#{rank}【相关图片】发现: {clean_p}')
            
            if not all_cleaned_parts:
                return {'success': False, 'error': '检索结果清洗后内容为空'}
            
            # 合并多条切片（用空行分隔）
            combined_content = '\n\n'.join(all_cleaned_parts)
            # 去重图片路径
            unique_images = list(dict.fromkeys(all_images))[:15]  # 最多15张
            
            logger.info(
                f'[KB] 合并完成: {len(all_cleaned_parts)}条文本切片, 总{len(combined_content)}字, '
                f'图片{len(unique_images)}张(已选{len(selected_chunks)}条, 2+3合并={merged_chunk3})'
            )

            return {
                'success': True,
                'content': combined_content,
                'images': unique_images,
                'count': len(candidate_chunks),
                'top_score': round(candidate_chunks[0].get('similarity', 0), 3),
                'keywords': keywords
            }
            
        except ImportError:
            logger.error('❌ 无法导入 requests 库')
            return {'success': False, 'error': '系统缺少 requests 库'}
        except Exception as e:
            logger.error(f'❌ 知识库检索异常: {e}')
            return {'success': False, 'error': str(e)}

    def _clean_chunk(self, raw: str) -> tuple:
        """
        清洗知识库原始切片，去除元数据格式，保留纯净中文叙述文字和图片引用。
        返回 (清洗后文本, 提取到的图片路径列表)
        """
        text = raw
        extracted_images = []

        # 1. 【章节】路径前缀删除
        text = re.sub(r'【章节】[^#\n]*#\s*', '', text)
        text = re.sub(r'【章节】[^\n]*\n?', '', text)

        # 2. 【相关图片】→ 提取图片路径而非删除（RagFlow 使用中文括号（））
        def _extract_images(match):
            nonlocal extracted_images
            img_block = match.group(0)
            # 同时匹配中文括号（）和英文括号()
            img_paths = re.findall(r'[（(]([^)）]*?\.(?:jpg|png|jpeg|gif|bmp))[）)]', img_block, re.IGNORECASE)
            for p in img_paths:
                clean_p = p.strip()
                if clean_p and clean_p not in extracted_images:
                    extracted_images.append(clean_p)
                    logger.info(f'[CleanChunk] Image: {clean_p}')
            return ''
        text = re.sub(r'【相关图片】.*?(?=\n\n|\Z)', _extract_images, text, flags=re.DOTALL)

        # 3. LaTeX formula
        text = re.sub(r'\$\$[^$]+\$\$', '', text)
        text = re.sub(r'\$[^$\n]{1,120}\$', '', text)

        # 4. (images/...) path -> extract but KEEP in text (支持中文括号)
        def _cap_img(m):
            nonlocal extracted_images
            p = m.group(1) if m.lastindex >= 1 else m.group(0)
            clean_p = p.strip()
            if clean_p and clean_p not in extracted_images:
                extracted_images.append(clean_p)
            return m.group(0)  # 保留原文，不删除
        text = re.sub(r'[（(](images/[^)）]+)[）)]', _cap_img, text, flags=re.IGNORECASE)

        # Keep hash filenames too (also keep in text)
        def _cap_hash(m):
            nonlocal extracted_images
            full = m.group(0)
            if full not in extracted_images:
                extracted_images.append(full)
            return full  # 保留原文
        text = re.sub(r'\b[a-f0-9]{20,}\.(jpg|png|jpeg)\b', _cap_hash, text)

        # 5. Section number at line start -> remove number keep text
        text = re.sub(r'^\s*\d+(\.\d+){1,3}\s+', '', text, flags=re.MULTILINE)

        # 6. Markdown # heading symbols
        text = re.sub(r'#\s*\d+(\.\d+){0,3}\s*[^#\n]{0,30}#\s*', ' ', text)
        text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)

        # 7. Figure reference labels
        text = re.sub(r'图\s*\d+[.\-]\d[-\d]*\s*', '', text)

        # 8. Empty parentheses leftover
        text = re.sub(r'[（(]\s*[）)]', '', text)

        # 9. Context after formula removal
        text = re.sub(r'(桩号|高程|面积|开采量|总量)\s+内[，,]', r'\1（数据略）内,', text)
        text = re.sub(r'(桩号|高程|面积|开采量|总量)\s*[，,。]', r'\1（数据略）,', text)
        text = re.sub(r'万\s+([，,。])', r'万（数据略）\1', text)
        text = re.sub(r'\s+([，,。；;])', r'\1', text)

        # 10. Compress whitespace
        text = re.sub(r'[ \t]{2,}', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text.strip()

        # 11. Truncate to reasonable length
        if len(text) > 600:
            truncated = text[:600].rsplit('。', 1)
            text = truncated[0] + '。' if len(truncated) > 1 else text[:600]

        return text, extracted_images

